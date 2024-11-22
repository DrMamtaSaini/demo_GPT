import streamlit as st
import openai
import pandas as pd
from fpdf import FPDF
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
import os
import json
from docx import Document
from docx.shared import Inches
from io import BytesIO
import requests
from PyPDF2 import PdfReader
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tempfile
from pathlib import Path
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
from docx import Document
from docx.shared import Inches
import tempfile

import streamlit as st
import openai
import json
import requests
import time
from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from fpdf import FPDF
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
import openai
from docx import Document
import re
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np


import streamlit as st
from fpdf import FPDF
from docx import Document
import openai

# Set OpenAI API key securely from Streamlit secrets
try:
    openai.api_key = st.secrets["api_key"]
except KeyError:
    st.error("API Key is missing in Streamlit secrets. Please add it and restart the app.")

# Load school credentials securely from Streamlit secrets
try:
    SCHOOL_CREDENTIALS = st.secrets["scho_credentials"]
except KeyError:
    st.error("School credentials are missing in Streamlit secrets. Please add them to proceed.")

# Load client configuration from JSON file with error handling
clients_config = {}
try:
    with open("clients_config.json") as config_file:
        clients_config = json.load(config_file)
except FileNotFoundError:
    st.warning("Configuration file 'clients_config.json' not found. Default settings will be used.")
except json.JSONDecodeError:
    st.error("Configuration file 'clients_config.json' is invalid. Please check the file format.")
except Exception as e:
    st.error(f"Unexpected error loading configuration: {e}")

# Helper function to retrieve client configuration based on client_id
def get_client_config(client_id):
    """Retrieves client configuration based on client_id. Returns None if not found."""
    return clients_config.get(client_id)

# Login page function with added error handling
def login_page():
    """Displays login page, authenticates users, and sets session states on successful login."""
    st.title("EduCreate Pro Login")
    
    # Collect username and password with placeholder values
    school_username = st.text_input("Username", placeholder="Enter username", key="username_input")
    school_password = st.text_input("Password", type="password", placeholder="Enter password", key="password_input")

    if st.button("Login", help="Click twice if needed to confirm login"):
        # Initialize login status to False, will update on successful authentication
        st.session_state['logged_in'] = False

        # Validate credentials against stored values in Streamlit secrets
        for school_id, credentials in SCHOOL_CREDENTIALS.items():
            if school_username == credentials["username"] and school_password == credentials["password"]:
                # Successful login: Set session state and retrieve client configuration
                st.session_state['logged_in'] = True
                st.session_state['client_id'] = school_id
                st.session_state['page'] = 'main'
                
                # Attempt to load client configuration after successful login
                client_config = get_client_config(school_id)
                if client_config:
                    st.session_state['client_config'] = client_config
                    st.success("Login successful!")
                else:
                    st.error("Client configuration not found. Please contact support.")
                break
        else:
            # Invalid credentials feedback
            st.error("Invalid credentials. Please try again or contact your administrator.")

        # Confirm if no valid credentials found in SCHOOL_CREDENTIALS
        if not st.session_state.get("logged_in"):
            st.warning("Please ensure your credentials are correct. Try again.")

def logout():
    """Handle user logout by resetting session state variables."""
    st.session_state['logged_in'] = False
    st.session_state['page'] = 'home'



def fetch_image(prompt, retries=5):
    """
    Fetches an image from the OpenAI API based on the given prompt.
    Includes retry logic for handling rate limits and network errors.
    
    Args:
        prompt (str): Text prompt for image generation.
        retries (int): Number of retry attempts for rate limit or network issues.
    
    Returns:
        BytesIO: In-memory image data if successful, None otherwise.
    """
    for attempt in range(retries):
        try:
            # Request image generation from OpenAI
            response = openai.Image.create(prompt=prompt, n=1, size="512x512")
            image_url = response['data'][0]['url']
            
            # Attempt to fetch the image from the URL provided by OpenAI
            image_response = requests.get(image_url)
            image_response.raise_for_status()  # Ensure the request was successful
            return BytesIO(image_response.content)  # Return image as BytesIO for easy handling
            
        except openai.error.RateLimitError:
            # Handle rate limit by retrying with exponential backoff
            if attempt < retries - 1:
                wait_time = 2 ** attempt  # Exponential backoff for retries
                st.warning(f"Rate limit reached. Retrying in {wait_time} seconds...")
                time.sleep(wait_time)
            else:
                st.error("Rate limit exceeded. Unable to fetch image after multiple attempts.")
                return None  # Return None to gracefully handle in the main app

        except requests.exceptions.RequestException as e:
            st.error(f"Network error while fetching image: {e}")
            return None  # Return None on network failure to allow error handling in the calling function

        except Exception as e:
            st.error(f"Unexpected error fetching image: {e}")
            return None  # Return None for unexpected errors
            
    # Fallback if all retries fail
    st.error("Failed to fetch image after multiple attempts.")
    return None




def generate_image_based_questions():
    """
    Generates image-based questions for quizzes and provides download options for the generated documents.
    """
    st.header("Generate Image-Based Questions")
    
    try:
        # Predefined options
        predefined_subjects = ["Science", "Geography", "Mathematics", "English", "History"]
        predefined_class_levels = ["Grade 1", "Grade 5", "Grade 10", "Grade 12"]
        predefined_max_marks = ["10","20", "50", "100"]
        predefined_durations = ["30 minutes", "1 hour", "2 hours"]

        # Dropdowns with add-new functionality
        subject = st.selectbox("Select Subject:", predefined_subjects + ["Add New"])
        if subject == "Add New":
            subject = st.text_input("Enter New Subject:")

        class_level = st.selectbox("Select Class Level:", predefined_class_levels + ["Add New"])
        if class_level == "Add New":
            class_level = st.text_input("Enter New Class Level:")

        max_marks = st.selectbox("Select Maximum Marks:", predefined_max_marks + ["Add New"])
        if max_marks == "Add New":
            max_marks = st.text_input("Enter New Maximum Marks:")

        duration = st.selectbox("Select Duration:", predefined_durations + ["Add New"])
        if duration == "Add New":
            duration = st.text_input("Enter New Duration:")

        # Other inputs
        topic = st.text_input("Select a Topic (e.g., Plants, Animals, Geography):", key="image_topic_input")
        num_questions = st.number_input("Enter the Number of Questions (minimum 5):", min_value=5, key="num_questions_input")
        question_type = st.selectbox("Choose Question Type:", ["MCQ", "True/False", "Yes/No"], key="question_type_select")

        # Generate Quiz Documents
        if st.button("Generate Quiz Document"):
            if not topic or not subject or not class_level or not max_marks or not duration:
                st.warning("Please fill in all required fields.")
                return

            try:
                # Temporary files for quiz documents
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file_without_answers, \
                     tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file_with_answers:

                    quiz_filename_without_answers = tmp_file_without_answers.name
                    quiz_filename_with_answers = tmp_file_with_answers.name

                    create_quiz_document(
                        topic, subject, class_level, max_marks, duration, num_questions, question_type,
                        include_answers=False, file_path=quiz_filename_without_answers
                    )
                    create_quiz_document(
                        topic, subject, class_level, max_marks, duration, num_questions, question_type,
                        include_answers=True, file_path=quiz_filename_with_answers
                    )

                    # Store filenames in session state for download
                    st.session_state["quiz_filename_without_answers"] = quiz_filename_without_answers
                    st.session_state["quiz_filename_with_answers"] = quiz_filename_with_answers

                st.success("Quiz documents generated successfully!")

            except Exception as e:
                st.error(f"Error generating quiz documents: {e}")

        # Download Buttons
        if "quiz_filename_without_answers" in st.session_state and "quiz_filename_with_answers" in st.session_state:
            try:
                with open(st.session_state["quiz_filename_without_answers"], "rb") as file:
                    st.download_button(
                        label="Download Quiz Document (without answers)",
                        data=file.read(),
                        file_name=Path(st.session_state["quiz_filename_without_answers"]).name
                    )

                with open(st.session_state["quiz_filename_with_answers"], "rb") as file:
                    st.download_button(
                        label="Download Quiz Document (with answers)",
                        data=file.read(),
                        file_name=Path(st.session_state["quiz_filename_with_answers"]).name
                    )

            except Exception as e:
                st.error(f"Error providing download options: {e}")
    
    except Exception as e:
        st.error(f"Error in Generate Image-Based Questions: {e}")

def create_quiz_document(topic, subject, class_level, max_marks, duration, num_questions, question_type, include_answers, file_path):
    """
    Creates a quiz document with questions and optional answers based on input parameters.
    
    Args:
        topic (str): Topic of the quiz.
        subject (str): Subject of the quiz.
        class_level (str): Grade or class level.
        max_marks (int): Maximum marks for the quiz.
        duration (str): Duration of the quiz.
        num_questions (int): Number of questions to include.
        question_type (str): Type of questions (e.g., MCQ, True/False).
        include_answers (bool): Whether to include answers in the document.
        file_path (str): File path to save the document.
    """
    try:
        document = Document()

        # Centered title and details
        document.add_heading('Assignment', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_heading(f'Grade: {class_level}', level=2).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_heading(f'Subject: {subject}', level=2).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_heading(f'Topic: {topic}', level=2).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph("\n")

        # Duration and Marks information
        details_paragraph = document.add_paragraph()
        details_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        details_paragraph.add_run(f'Duration: {duration}    Max. Marks: {max_marks}').bold = True

        # Initialize list to store correct answers if required
        answers = []
        
        for i in range(num_questions):
            # Generate question text and options
            subtopic = f"subtopic_{i % 3 + 1}"  # Placeholder subtopics
            question_text, options, correct_answer = generate_question_and_options(topic, class_level, question_type, subtopic)
            document.add_paragraph(f'Q{i+1}: {question_text}')
            
            # Add each option
            for option in options:
                document.add_paragraph(option)
                
            # Save correct answer if including answers
            if include_answers:
                answers.append(f'Q{i+1}: {correct_answer}')

            # Fetch and add an image if applicable
            image_prompt = f"Image of {subtopic} related to {topic}"
            image_data = fetch_image(image_prompt)
            if image_data:
                document.add_picture(image_data, width=Inches(2))
                
            document.add_paragraph("\n")  # Space between questions

        # Add answers at the end if include_answers is set to True
        if include_answers:
            document.add_paragraph("\nAnswers:\n")
            for answer in answers:
                document.add_paragraph(answer)
        else:
            document.add_paragraph("\nAnswers:\n")
            for i in range(num_questions):
                document.add_paragraph(f'Q{i+1}: ________________')

        # Save document with error handling
        document.save(file_path)
        st.success(f"Quiz document saved successfully as {file_path}")
        
    except Exception as e:
        st.error(f"Error generating quiz document: {e}")


# Function to generate a PDF file for reports



def generate_pdf(content, title, file_name):
    """
    Generates a PDF report with the provided content.
    
    Args:
        content (str): The text content to include in the PDF.
        title (str): The title for the PDF report.
        file_name (str): The file path to save the PDF.
    """
    try:
        pdf = FPDF()
        pdf.add_page()
        
        # Title with encoding handling
        pdf.set_font("Arial", "B", 16)
        title_sanitized = sanitize_text(title)
        pdf.cell(0, 10, title_sanitized, ln=True, align='C')
        pdf.ln(10)
        
        # Content with encoding handling
        pdf.set_font("Arial", size=12)
        for line in content.split('\n'):
            pdf.cell(0, 10, sanitize_text(line), ln=True)
        
        # Save the PDF file
        pdf.output(file_name)
        st.success(f"PDF generated successfully and saved as {file_name}")
        
    except Exception as e:
        st.error(f"Error generating PDF: {e}")


# Function to send an email with PDF attachment

def send_email_with_pdf(to_email, subject, body, file_name):
    """
    Sends an email with a PDF attachment to the specified recipient.
    
    Args:
        to_email (str): Recipient's email address.
        subject (str): Email subject.
        body (str): Email body.
        file_name (str): Path to the PDF file to attach.
    """
    from_email = st.secrets.get("email")
    password = st.secrets.get("app_password")
    
    # Error handling for missing credentials
    if not from_email or not password:
        st.error("Email credentials are missing. Check Streamlit secrets.")
        return
    
    # Set up email message
    msg = MIMEMultipart()
    msg['From'] = from_email
    msg['To'] = to_email
    msg['Subject'] = subject
    msg.attach(MIMEText(body, 'plain'))
    
    # Attach PDF file with error handling
    try:
        with open(file_name, "rb") as attachment:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(attachment.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f'attachment; filename= {file_name}')
            msg.attach(part)
    except FileNotFoundError:
        st.error(f"Attachment file {file_name} not found. Please check the file path.")
        return
    except Exception as e:
        st.error(f"Error attaching file {file_name}: {e}")
        return
    
    # Send email with error handling
    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(from_email, password)
        server.sendmail(from_email, to_email, msg.as_string())
        server.quit()
        st.success(f"Email sent to {to_email} with the attached PDF report!")
    except smtplib.SMTPException as e:
        st.error(f"SMTP error occurred: {e}")
    except Exception as e:
        st.error(f"Unexpected error sending email: {e}")



# Function to generate a quiz document
def generate_question_and_options(topic, class_level, question_type, subtopic):
    """
    Generates a question and options using OpenAI's API based on input parameters.
    
    Args:
        topic (str): Main topic for the question.
        class_level (str): Grade or class level.
        question_type (str): Type of question (e.g., MCQ).
        subtopic (str): Specific subtopic within the main topic.

    Returns:
        tuple: Question text, options, and correct answer.
    """
    prompt = (
        f"Create a {question_type} question about {subtopic} for a {class_level} level quiz on {topic}. "
        f"Include four answer choices labeled A, B, C, and D, and indicate the correct answer."
    )
    
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}]
        )
        full_text = response['choices'][0]['message']['content'].strip()
        lines = full_text.split("\n")
        question = lines[0]
        options = lines[1:5]  # First four lines after question are options
        correct_answer = lines[5] if len(lines) > 5 else "Correct Answer Not Provided"
        
        if not question or not options:
            raise ValueError("Incomplete question or options received.")
        
        return question, options, correct_answer
    
    except openai.error.OpenAIError as e:
        st.error(f"OpenAI API error: {e}")
        return "Error generating question.", ["A) Error", "B) Error", "C) Error", "D) Error"], "Answer unavailable"
    except Exception as e:
        st.error(f"Unexpected error generating question: {e}")
        return "Error generating question.", ["A) Error", "B) Error", "C) Error", "D) Error"], "Answer unavailable"


def save_contentt_as_doc(content, file_name_docx, image_data=None, single_image=True):
    """
    Saves content as a DOCX file with optional images.
    
    Args:
        content (str): Text content to include in the document.
        file_name_docx (str): File path for saving the document.
        image_data (BytesIO): Optional image data to embed.
        single_image (bool): Whether to add a single image at the top or one per question.
    """
    try:
        document = Document()
        
        # If single_image is True, add one image at the top; if False, add one per question
        if single_image and image_data:
            document.add_picture(image_data, width=Inches(2))
            document.add_paragraph("\n")

        for line in content.split("\n"):
            document.add_paragraph(line)
            if not single_image and image_data:
                document.add_picture(image_data, width=Inches(2))
                document.add_paragraph("\n")
        
        document.save(file_name_docx)
        st.success(f"Document saved successfully as {file_name_docx}")

    except Exception as e:
        st.error(f"Error saving document {file_name_docx}: {e}")

def extract_weak_topics(assessment_content):
    """
    Identifies weak topics based on assessment content using OpenAI's API.
    
    Args:
        assessment_content (str): The assessment report content to analyze.

    Returns:
        list: Weak topics if found, else an empty list.
    """
    prompt = f"""
    Analyze the following assessment report content. Identify and list all topics and subtopics where 'Concept Clarity' is marked as 'No'.
    
    Assessment Content:
    {assessment_content}
    """
    
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}]
        )
        weak_topics = response['choices'][0]['message']['content']
        if not weak_topics:
            raise ValueError("No weak topics identified.")
        return weak_topics.split("\n")

    except openai.error.OpenAIError as e:
        st.error(f"OpenAI API error while extracting weak topics: {e}")
        return []
    except Exception as e:
        st.error(f"Unexpected error extracting weak topics: {e}")
        return []

def generate_personalized_material(weak_topics):
    """
    Generates personalized learning material based on weak topics.
    
    Args:
        weak_topics (list): List of weak topics to address.

    Returns:
        str: Learning material content.
    """
    prompt = f"Create learning material covering the following topics: {', '.join(weak_topics)}."
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}]
        )
        material = response['choices'][0]['message']['content']
        return material if material else "No material generated."

    except openai.error.OpenAIError as e:
        st.error(f"OpenAI API error generating material: {e}")
        return "Error generating material."
    except Exception as e:
        st.error(f"Unexpected error generating material: {e}")
        return "Error generating material."


def send_email_with_attachments(to_email, subject, body, attachments):
    """
    Sends an email with multiple attachments to the specified recipient.
    
    Args:
        to_email (str): Recipient's email address.
        subject (str): Email subject.
        body (str): Email body.
        attachments (list): List of file paths for attachments.
    """
    from_email = st.secrets.get("email")
    password = st.secrets.get("app_password")
    
    if not from_email or not password:
        st.error("Email credentials missing. Check Streamlit secrets.")
        return
    
    try:
        msg = MIMEMultipart()
        msg['From'] = from_email
        msg['To'] = to_email
        msg['Subject'] = subject
        msg.attach(MIMEText(body, 'plain'))
        
        for file_path in attachments:
            try:
                with open(file_path, "rb") as attachment:
                    part = MIMEBase("application", "octet-stream")
                    part.set_payload(attachment.read())
                    encoders.encode_base64(part)
                    part.add_header("Content-Disposition", f"attachment; filename= {file_path}")
                    msg.attach(part)
            except FileNotFoundError:
                st.error(f"Attachment file {file_path} not found.")
                return
            except Exception as e:
                st.error(f"Error attaching file {file_path}: {e}")
                return
        
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(from_email, password)
        server.sendmail(from_email, to_email, msg.as_string())
        server.quit()
        st.success(f"Email sent to {to_email} with attachments.")
    
    except smtplib.SMTPException as e:
        st.error(f"SMTP error: {e}")
    except Exception as e:
        st.error(f"Unexpected error sending email: {e}")










# Function to generate question using GPT based on input
def generate_question(topic, class_level, question_type, subtopic):
    try:
        prompt = f"Generate a {question_type} question on the topic '{topic}' for {class_level} on the subtopic '{subtopic}'. Include a question text and answer options."
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}]
        )
        return response['choices'][0]['message']['content']
    except Exception as e:
        print(f"Error in generating question: {e}")
        return "Error generating question."

# Function to create a quiz document with enhanced error handling for images
def create_quiz1_document(topic, subject, class_level, max_marks, duration, num_questions, question_type, include_answers, include_images, file_path):
    document = Document()
    try:
        # Centered main title
        document.add_heading('Quiz', level=1).alignment = 1  # Center align
        document.add_heading(f'Class: {class_level}', level=2).alignment = 1
        document.add_heading(f'Subject: {subject}', level=2).alignment = 1
        document.add_heading(f'Topic: {topic}', level=2).alignment = 1
        document.add_paragraph("\n")  # Blank line for spacing

        # Add Duration and Max Marks on the same line, centered
        details_paragraph = document.add_paragraph()
        details_paragraph.alignment = 1  # Center align
        details_paragraph.add_run(f'Duration: {duration}').bold = True
        details_paragraph.add_run("    ")
        details_paragraph.add_run(f'Max. Marks: {max_marks}').bold = True
        document.add_paragraph("\n")  # Extra line for spacing

        # Loop to add questions and images
        for i in range(num_questions):
            question_text = f"Sample question {i+1} about {topic}."
            document.add_paragraph(f'Q{i+1}: {question_text}')

            # Conditionally add image if include_images is True
            if include_images:
                image_prompt = f"Image of {topic} for {class_level} related to {subject}"
                image_data = fetch_image(image_prompt)
                if image_data:
                    document.add_picture(image_data, width=Inches(2))
            
            # Add answer options based on question type
            if question_type == "MCQ":
                options = ["A) Option 1", "B) Option 2", "C) Option 3", "D) Option 4"]
                for option in options:
                    document.add_paragraph(option)
            elif question_type == "true/false":
                document.add_paragraph("A) True")
                document.add_paragraph("B) False")
            elif question_type == "yes/no":
                document.add_paragraph("A) Yes")
                document.add_paragraph("B) No")

            # Add correct answer if include_answers is True
            if include_answers:
                correct_answer = "B) Sample Answer"
                document.add_paragraph(f"Answer: {correct_answer}")

            document.add_paragraph("\n")  # Add spacing after each question

        # Add answer section if answers are not included inline
        if not include_answers:
            document.add_paragraph("\nAnswers:\n")
            for i in range(num_questions):
                document.add_paragraph(f'Q{i+1}: ________________')

        # Save the document to the specified file path
        document.save(file_path)
    except Exception as e:
        print(f"Error in creating quiz document: {e}")

# Function to generate an assessment report as PDF
def generatereport_pdf(content, title, filename, student_name, student_id, assessment_id, exam_type, subject):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    try:
        # Title - Assessment Report
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, "Assessment Report", ln=True, align="C")

        # Centered Exam Type and Subject
        pdf.set_font("Arial", "", 12)
        pdf.cell(0, 10, f"Exam Type: {exam_type}", ln=True, align="C")
        pdf.cell(0, 10, f"Subject: {subject}", ln=True, align="C")
        pdf.ln(10)  # Blank line for spacing

        # Detailed Summary Heading
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 10, "Detailed Summary Report", ln=True, align="L")
        pdf.ln(5)

        # Student information
        pdf.set_font("Arial", "", 12)
        pdf.cell(0, 10, f"Student Name: {student_name}", ln=True)
        pdf.cell(0, 10, f"Student ID: {student_id}", ln=True)
        pdf.cell(0, 10, f"Assessment ID: {assessment_id}", ln=True)
        pdf.ln(10)

        # Add the report content
        pdf.multi_cell(0, 10, content)
        
        # Save the PDF
        pdf.output(filename)
    except Exception as e:
        print(f"Error in generating report PDF: {e}")

# Function to read DOCX content with debug statements
def read_docx(file):
    try:
        doc = Document(file)
        full_text = "\n".join([repr(para.text) for para in doc.paragraphs])
        print("DEBUG - Full Document Content with Hidden Characters:\n", full_text)  # Show all text details
        return full_text
    except Exception as e:
        print(f"Error reading DOCX file: {e}")
        return "Error reading file."

# Function to generate personalized assignments
def generate_personalized_assignment(weak_topics):
    try:
        prompt = f"Create an assignment based on the following topics for practice: {', '.join(weak_topics)}. Include questions that reinforce the concepts."
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}]
        )
        return response['choices'][0]['message']['content']
    except Exception as e:
        print(f"Error generating personalized assignment: {e}")
        return "Error generating assignment."

# Function to generate a PDF file with enhanced encoding
def generate_pdf(content, title, file_name):
    pdf = FPDF()
    pdf.add_page()

    try:
        # Set title with error handling for encoding
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, title.encode("latin-1", "replace").decode("latin-1"), ln=True, align='C')
        pdf.ln(10)

        # Add content with encoding error handling
        pdf.set_font("Arial", size=12)
        for line in content.split('\n'):
            sanitized_line = line.encode("latin-1", "replace").decode("latin-1")
            pdf.cell(200, 10, sanitized_line, ln=True)

        # Output PDF to file
        pdf.output(file_name)
    except Exception as e:
        print(f"Error in generating PDF: {e}")


# Function to generate a PDF file with enhanced encoding
def generatenew_pdf(content, title, file_name):
    """
    Generates a PDF report with sanitized text to handle Unicode issues.

    Args:
        content (str): The text content to include in the PDF.
        title (str): The title for the PDF report.
        file_name (str): The file path to save the PDF.
    """
    pdf = FPDF()
    pdf.add_page()

    # Set title
    pdf.set_font("Arial", "B", 16)
    sanitized_title = sanitize_text_for_pdf(title)
    pdf.cell(0, 10, sanitized_title, ln=True, align='C')
    pdf.ln(10)

    # Add content
    pdf.set_font("Arial", size=12)
    sanitized_content = sanitize_text_for_pdf(content)
    for line in sanitized_content.split("\n"):
        pdf.cell(0, 10, line, ln=True)

    # Save the PDF file
    try:
        pdf.output(file_name)
        print(f"PDF saved as {file_name}")
    except Exception as e:
        print(f"Error in saving PDF: {e}")


# Function to save content to a DOCX file with error handling
def save_content_as_doc(content, file_name):
    try:
        doc = Document()
        for line in content.split("\n"):
            doc.add_paragraph(line)
        doc.save(file_name)
    except Exception as e:
        print(f"Error saving content as DOCX: {e}")





# Function to sanitize text by replacing unsupported characters
def sanitize_text(text):
    return text.encode('latin-1', 'replace').decode('latin-1')


def create_lesson_plan():
    """
    Collects input to generate a lesson plan and provides feedback for successful or unsuccessful generation.
    """
    st.header("Lesson Plan Creation")

    try:
        # Predefined options for subjects, boards, grades, and durations
        predefined_subjects = ["Mathematics", "Science", "English", "History", "Geography"]
        predefined_boards = ["CBSE", "ICSE", "State Board", "IB", "Cambridge"]
        predefined_grades = ["Grade 1", "Grade 5", "Grade 10", "Grade 12"]
        predefined_durations = ["30 minutes", "45 minutes", "1 hour", "1.5 hours"]

        # Comprehensive list of languages with Kiswahili included
        mediums = [
            "English", "Hindi", "Gujarati", "Marathi", "Tamil", "Telugu", "Bengali", 
            "Japanese", "Chinese (Mandarin)", "Korean", "French", "German", "Spanish", 
            "Italian", "Russian", "Arabic", "Portuguese", "Kiswahili", "Malay", "Thai", 
            "Vietnamese", "Urdu", "Punjabi", "Persian", "Turkish", "Dutch", "Greek", "Add New"
        ]

        # Subject selection with option to add a custom subject
        subject = st.selectbox("Select Subject:", predefined_subjects + ["Add New"])
        if subject == "Add New":
            subject = st.text_input("Enter New Subject:", key="subject_input")

        # Grade selection with option to add a custom grade
        grade = st.selectbox("Select Grade/Class:", predefined_grades + ["Add New"])
        if grade == "Add New":
            grade = st.text_input("Enter New Grade/Class:", key="grade_input")

        # Board selection with option to add a custom board
        board = st.selectbox("Select Education Board:", predefined_boards + ["Add New"])
        if board == "Add New":
            board = st.text_input("Enter New Education Board:", key="board_input")

        # Medium selection with English as the default
        medium = st.selectbox("Select Medium:", mediums, index=0)
        if medium == "Add New":
            medium = st.text_input("Enter New Medium:")

        # Lesson Duration selection with option to add a custom duration
        duration = st.selectbox("Select Lesson Duration:", predefined_durations + ["Add New"])
        if duration == "Add New":
            duration = st.text_input("Enter New Duration (e.g., 2 hours):", key="duration_input")

        # Topic input (removed dropdown for predefined topics)
        topic = st.text_input("Enter Lesson Topic:", key="topic_input")

        # Generate lesson plan
        if st.button("Generate Lesson Plan"):
            if not subject or not grade or not board or not duration or not topic or not medium:
                st.warning("Please fill in all required fields.")
                return

            try:
                # Call the lesson plan generation function
                lesson_plan = generate_lesson_plan(subject, grade, board, duration, topic, medium)
                st.write("### Generated Lesson Plan")
                st.write(lesson_plan)

                # Save lesson plan as DOCX and PDF, with error handling
                file_name_docx = f"{subject}_{grade}_LessonPlan.docx"
                save_content_as_doc(lesson_plan, file_name_docx)
                with open(file_name_docx, "rb") as file:
                    st.download_button(label="Download Lesson Plan as DOCX", data=file.read(), file_name=file_name_docx)

                file_name_pdf = f"{subject}_{grade}_LessonPlan.pdf"
                generate_pdf(lesson_plan, f"Lesson Plan: {subject} -  {grade}", file_name_pdf)
                with open(file_name_pdf, "rb") as file:
                    st.download_button(label="Download Lesson Plan as PDF", data=file.read(), file_name=file_name_pdf)

            except Exception as e:
                st.error(f"Error generating lesson plan or saving files: {e}")

    except Exception as e:
        st.error(f"Error in lesson plan creation: {e}")


def generate_lesson_plan(subject, grade, board, duration, topic, medium):
    """
    Generates a comprehensive lesson plan based on the given inputs.
    """
    prompt = f"""
    Create a comprehensive lesson plan for teaching {subject} to {grade} under the {board} board. 
    The lesson medium is {medium}. The lesson duration is {duration}, and the topic of the lesson is {topic}. 
    The lesson plan should include:

    - Lesson Title and Duration
    - Learning Objectives
    - Materials and Resources Needed
    - Detailed Lesson Flow:
        - Introduction (5-10 minutes)
        - Core Teaching Segment with explanations and examples
        - Interactive Activities for engagement
        - Assessment and Recap to measure understanding
        - Homework/Assignments for reinforcement
    - Date and Schedule field

    Ensure flexibility to adapt to different student needs, learning speeds, and teaching styles.
    """
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "system", "content": prompt}]
    )
    return response['choices'][0]['message']['content']


# Main function
import re

def validate_email(email):
    """Validate email format."""
    return re.match(r"[^@]+@[^@]+\.[^@]+", email) is not None

def sanitize_text_input(text):
    """Sanitize text input to prevent special character injection."""
    return re.sub(r"[^\w\s]", "", text)  # Removes special characters except spaces and alphanumeric

def validate_student_id(student_id):
    """Validate student ID format (only alphanumeric allowed)."""
    return re.match(r"^[a-zA-Z0-9]+$", student_id) is not None


def show_home():
    """
    Displays the home page with custom-styled option cards for each section.
    """
    try:
        st.markdown("""
            <div style='text-align: center; font-size: 18px; color: #4B0082; padding: 20px 0;'>
                Welcome to your all-in-one platform for creating educational content, lesson plans, and student assessments.
            </div>
        """, unsafe_allow_html=True)

        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("""
            <div class="option-card">
                <h3 style='color: #4B0082; text-align: center;'>Content Creator</h3>
                <p style='text-align: center;'>Generate quizzes, sample papers, and assignments for students.</p>
            </div>
            """, unsafe_allow_html=True)
        
        with col2:
            st.markdown("""
            <div class="option-card">
                <h3 style='color: #4B0082; text-align: center;'>Lesson Planner</h3>
                <p style='text-align: center;'>Create detailed lesson plans with objectives and resources.</p>
            </div>
            """, unsafe_allow_html=True)
        
        col3, col4, col5 = st.columns(3)
        
        with col3:
            st.markdown("""
            <div class="option-card">
                <h3 style='color: #4B0082; text-align: center;'>Assessment Assistant</h3>
                <p style='text-align: center;'>Generate comprehensive student assessments and progress reports.</p>
            </div>
            """, unsafe_allow_html=True)
        
        with col4:
            st.markdown("""
            <div class="option-card">
                <h3 style='color: #4B0082; text-align: center;'>Image-Based Question Generator</h3>
                <p style='text-align: center;'>Generate image-based quizzes (MCQ, True/False, Yes/No).</p>
            </div>
            """, unsafe_allow_html=True)
        
        with col5:
            st.markdown("""
            <div class="option-card">
                <h3 style='color: #4B0082; text-align: center;'>Analyze Report and Generate Graph</h3>
                <p style='text-align: center;'>Analyze Report and Generate Graph</p>
            </div>
            """, unsafe_allow_html=True)

        st.markdown("<div class='center' style='padding: 20px;'>", unsafe_allow_html=True)
        if st.button("Get Started Today"):
            st.session_state['task'] = "Create Educational Content"
        st.markdown("</div>", unsafe_allow_html=True)
    
    except Exception as e:
        st.error(f"Error displaying home page: {e}")

def create_educational_content():
    """
    Collects input to generate a question paper and handles all errors gracefully.
    """
    st.header("Question Paper Creation")
    
    try:
        # Predefined options for dropdowns and checkboxes
        BOARDS = ["CBSE", "ICSE", "State Board", "IB", "Cambridge", "Others"]
        GRADES = [f"Class {i}" for i in range(1, 13)] + ["Others"]
        SUBJECTS = ["Mathematics", "Science", "History", "Geography", "English", "Others"]
        DIFFICULTY_LEVELS = ["Easy", "Medium", "Hard"]
        CATEGORIES = [
            "Value-based Questions", "Competency Questions", "Mixed of your choice"
        ]
        QUESTION_TYPES = [
            "True/False", "Yes/No", "MCQs", 
            "Very Short Answers", "Short Answers", 
            "Long Answers", "Very Long Answers"
        ]
        MEDIUMS = ["English", "Hindi", "Tamil", "Bengali", "Kiswahili", "Others"]

        # Dropdowns with "Others" option for custom input
        board = st.selectbox("Select Education Board:", BOARDS)
        if board == "Others":
            board = st.text_input("Enter Board Name:")

        standard = st.selectbox("Select Standard/Class:", GRADES)
        if standard == "Others":
            standard = st.text_input("Enter Standard/Class:")

        subject = st.selectbox("Select Subject:", SUBJECTS)
        if subject == "Others":
            subject = st.text_input("Enter Subject Name:")

        medium = st.selectbox("Select Medium:", MEDIUMS)
        if medium == "Others":
            medium = st.text_input("Enter Medium of Instruction:")

        topics = st.text_area("Enter Topics (comma-separated):", placeholder="E.g., Algebra, Photosynthesis")

        total_marks = st.number_input("Enter Total Marks (optional):", min_value=10, step=5)

        time_duration = st.text_input("Enter Time Duration (e.g., 60 minutes, optional):")

        question_types = st.multiselect("Select Question Types (optional):", QUESTION_TYPES)

        difficulty = st.selectbox("Select Difficulty Level:", DIFFICULTY_LEVELS)

        # Category selection with checkboxes
        st.markdown("### Select Question Categories")
        selected_categories = []
        for category in CATEGORIES:
            if st.checkbox(category, key=category):
                selected_categories.append(category)

        # New field: Number of Questions
        num_questions = st.number_input("Enter Number of Questions:", min_value=1, max_value=50, step=1, value=10)

        include_solutions = st.radio("Would you like to include solutions?", ["Yes", "No"]) == "Yes"

        # Generate question paper button
        if st.button("Generate Question Paper"):
            if not board or not standard or not subject or not topics:
                st.warning("Please fill in all required fields: Board, Standard, Subject, and Topics.")
                return
            
            with st.spinner("Generating question paper..."):
                # Generate content using OpenAI
                content = generate_content(
                    board, standard, topics, total_marks, 
                    time_duration, question_types, difficulty, selected_categories, 
                    include_solutions, num_questions
                )
                
                # Display content
                st.write("### Generated Question Paper")
                st.write(content)

                # Save and allow download as DOCX and PDF
                try:
                    docx_file = create_answer_sheet_docx(content)

                    st.download_button(
                        label="Download Question Paper as DOCX",
                        data=docx_file,
                        file_name=f"Question_Paper_{standard}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                except Exception as e:
                    st.error(f"Error saving or downloading files: {e}")

    except Exception as e:
        st.error(f"Error in question paper creation: {e}")


def generate_content(board, standard, topics, total_marks, time_duration, question_types, difficulty, selected_categories, include_solutions, num_questions):
    """
    Generates a question paper based on the inputs.
    """
    category_list = ", ".join(selected_categories)
    
    # General content generation
    prompt = f"""
    You are an educational content creator. Create a question paper for the {board} board, {standard} class. 
    Based on the topics: {topics}. The question paper should contain {num_questions} questions, of {total_marks} marks 
    and a time duration of {time_duration}. The question types should include {', '.join(question_types)}, with 
    a difficulty level of {difficulty}. The categories of questions should be: {category_list}.
    """
    
    if include_solutions:
        prompt += " Include the solution set."
    else:
        prompt += " Only include the question set without solutions."
    
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "system", "content": prompt}]
        )
        return response['choices'][0]['message']['content']
    except Exception as e:
        raise Exception(f"Error generating question paper: {e}")




import matplotlib.pyplot as plt
from PyPDF2 import PdfReader

from fpdf import FPDF

def parse_pdf_with_openai(file):
    """
    Extracts text from the PDF file and uses OpenAI to parse topics and concept clarity.

    Args:
        file: Uploaded PDF file.
    
    Returns:
        dict: Parsed topics with their clarity percentages.
    """
    try:
        # Step 1: Extract text from the PDF
        reader = PdfReader(file)
        content = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                content += page_text

        # Debugging: Display extracted content
       # st.write("### Extracted Content")
        #st.text(content)

        # Step 2: Define the OpenAI prompt
        prompt = f"""
        The following text is an assessment report. Extract the topics and their concept clarity. 
        If a topic is marked as 'Yes' for concept clarity, set the percentage to 100. If marked as 'No', set it to 0. 
        Return the results as a JSON object with the format:
        {{
            "Topic 1": percentage,
            "Topic 2": percentage,
            ...
        }}
        
        Assessment Report:
        {content}
        """

        # Step 3: Use OpenAI to parse the report
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}]
        )
        parsed_data = response['choices'][0]['message']['content'].strip()

        # Step 4: Convert the response to a dictionary
        try:
            topics_data = json.loads(parsed_data)
            return topics_data
        except json.JSONDecodeError:
            st.error("Failed to parse the response into JSON format.")
            return {}

    except Exception as e:
        st.error(f"Error during parsing with OpenAI: {e}")
        return {}



import matplotlib.pyplot as plt

def generate_clarity_graph(data):
    """
    Generates a vertical bar chart for topics with concept clarity percentages.

    Args:
        data (dict): A dictionary with topics as keys and clarity percentages as values.
    """
    if not data:
        st.error("No data available to generate the graph.")
        return

    # Extract topics and percentages
    topics = list(data.keys())
    clarity_percentages = list(data.values())

    # Assign colors based on clarity zones
    colors = []
    for clarity in clarity_percentages:
        if clarity <= 30:
            colors.append("red")
        elif clarity <= 50:
            colors.append("orange")
        elif clarity <= 74:
            colors.append("yellow")
        else:
            colors.append("green")

    # Plot the vertical bar graph
    plt.figure(figsize=(12, 6))
    plt.bar(topics, clarity_percentages, color=colors, edgecolor="black")
    plt.xticks(rotation=45, ha='right', fontsize=10)  # Rotate topic labels for better readability
    plt.xlabel("Topics", fontsize=14)
    plt.ylabel("Clarity Percentage (%)", fontsize=14)
    plt.title("Concept Clarity by Topic", fontsize=16, fontweight="bold")
    plt.ylim(0, 100)  # Set y-axis limits
    plt.tight_layout()

    # Display the graph in Streamlit
    st.pyplot(plt)

def report_analysis_with_openai():
    st.header("Analyze Report Using OpenAI")
    
    uploaded_file = st.file_uploader("Upload Assessment Report (PDF)", type=["pdf"])
    
    if uploaded_file:
        try:
            # Parse the report using OpenAI
            topics_data = parse_pdf_with_openai(uploaded_file)

            # Display the parsed data
            if topics_data:
               # st.write("### Parsed Topics and Clarity")
                #st.json(topics_data)

                # Generate a graph for visualization
                st.write("### Concept Clarity Graph")
                generate_clarity_graph(topics_data)
            else:
                st.warning("No topics or concept clarity data found in the report.")
        
        except Exception as e:
            st.error(f"Error analyzing the report: {e}")




def sanitize_text_for_pdf(text):
    """Sanitizes text to ensure compatibility with PDF generation."""
    return text.encode('latin-1', 'replace').decode('latin-1')

# Example Usage
    pdf.cell(0, 10, sanitize_text_for_pdf(line), ln=True)







def student_assessment_assistant():
    """
    Collects inputs to generate student assessment reports, handles errors in report generation,
    and provides download and email options.
    """
    st.header("Student Assessment Assistant")

    try:
        # Collect student information with unique labels for each field
        student_name = st.text_input("Enter Student Name", key="student_name_input")
        student_id = st.text_input("Enter Student ID", key="student_id_input")
        assessment_id = st.text_input("Enter Assessment ID", key="assessment_id_input")
        class_name = st.text_input("Enter Class/Standard", key="class_name_input")
        email_id = st.text_input("Enter Parent's Email ID", key="email_id_input")
        exam_type = st.text_input("Enter Exam Type (e.g., Midterm, Final Exam)", key="exam_type_input")
        subject = st.text_input("Enter Subject", key="subject_input")

        # Upload files with DOCX format
        question_paper = st.file_uploader("Upload Question Paper (DOCX)", type=["docx"], key="question_paper_uploader")
        marking_scheme = st.file_uploader("Upload Marking Scheme (DOCX)", type=["docx"], key="marking_scheme_uploader")
        answer_sheet = st.file_uploader("Upload Student's Answer Sheet (DOCX)", type=["docx"], key="answer_sheet_uploader")

        if st.button("Generate and Send Reports"):
            if student_id and assessment_id and email_id and question_paper and marking_scheme and answer_sheet:
                try:
                    # Read DOCX files
                    question_paper_content = read_docx(question_paper)
                    marking_scheme_content = read_docx(marking_scheme)
                    answer_sheet_content = read_docx(answer_sheet)
                    
                    # Generate assessment report prompt
                    prompt = f"""
You are an educational assessment assistant. Using the question paper, marking scheme, and answer sheet, evaluate the student's answers.
Avoid using any special characters or bullet points for emphasis. Present the report in a clear, concise manner suitable for parents and teachers.
Please generate a detailed assessment report in the following format:

1. Question Analysis - For each question:
    - Topic
    - Subtopic
    - Question Number
    - Score based on answer accuracy and relevance
    - Concept Clarity (Yes/No)
    - Feedback and Suggestions
    - Solution if concept clarity is 'NO'

2. Summary Report - Include:
    - Final Score
    - Grade
    - Areas of Strength
    - Areas for Improvement
    - Final Remarks

Information for Report:
Student Name: {student_name}
Student ID: {student_id}
Class: {class_name}
Assessment ID: {assessment_id}
Exam Type: {exam_type}
Subject: {subject}

Question Paper:
{question_paper_content}

Marking Scheme:
{marking_scheme_content}

Answer Sheet:
{answer_sheet_content}
"""

                    response = openai.ChatCompletion.create(
                        model="gpt-3.5-turbo",
                        messages=[{"role": "system", "content": prompt}]
                    )
                    report = response['choices'][0]['message']['content']
                    st.write(" Assessment Report")
                    st.write(report)


                    
                    # Extract weak topics from report
                    weak_topics_prompt = f"Identify topics and subtopics where 'Concept Clarity' is marked as 'No'.\n\nAssessment Report:\n{report}"
                    weak_response = openai.ChatCompletion.create(
                        model="gpt-3.5-turbo",
                        messages=[{"role": "user", "content": weak_topics_prompt}]
                    )
                    weak_topics = weak_response['choices'][0]['message']['content'].strip().split("\n")
                                     
                    # Clean and format the weak topics into plain text
                    formatted_weak_topics = "Weak Topics Identified:\n"
                    for line in weak_topics:
                        if line.strip():  # Exclude empty lines
                            formatted_weak_topics += line.strip() + "\n"

                    # Display weak topics in plain text format
                    st.write("### Weak Topics")
                    st.text(formatted_weak_topics)  # Use st.text for plain text display

                    

                    # Generate learning material and assignment
                    learning_material_prompt = f"Create personalized learning material covering the following weak areas: {', '.join(weak_topics)}. Include explanations, examples, and practice questions."
                    assignment_prompt = f"Create an assignment based on the following weak areas for the student to improve: {', '.join(weak_topics)}. Include questions that reinforce concepts with solutions."

                    learning_material_response = openai.ChatCompletion.create(
                        model="gpt-3.5-turbo",
                        messages=[{"role": "user", "content": learning_material_prompt}]
                    )
                    assignment_response = openai.ChatCompletion.create(
                        model="gpt-3.5-turbo",
                        messages=[{"role": "user", "content": assignment_prompt}]
                    )

                    learning_material = learning_material_response['choices'][0]['message']['content']
                    assignment = assignment_response['choices'][0]['message']['content']
                    st.write("Personalized Learning Material")
                    st.write(learning_material)
                    st.write("Practice Assignment")
                    st.write(assignment)

                    # Generate PDFs
                    assessment_report_pdf = f"assessment_report_{student_id}.pdf"
                    generatereport_pdf(report, "Assessment Report", assessment_report_pdf, student_name, student_id, assessment_id, exam_type, subject)

                    learning_material_pdf = f"learning_material_{student_id}.pdf"
                    generate_pdf(learning_material, "Personalized Learning Material", learning_material_pdf)

                    assignment_pdf = f"assignment_{student_id}.pdf"
                    generate_pdf(assignment, "Personalized Assignment", assignment_pdf)

                    # Store file paths for download
                    st.session_state['assessment_report_pdf'] = assessment_report_pdf
                    st.session_state['learning_material_pdf'] = learning_material_pdf
                    st.session_state['assignment_pdf'] = assignment_pdf

                    st.success("Reports generated successfully and are ready for download.")

                    # Email the reports to the parent
                    subject = f"Assessment Reports for {student_name}"
                    body = f"""
Dear Parent,

Please find attached the assessment reports for {student_name}:

1. *Assessment Report*: Detailed evaluation of {student_name}'s performance.
2. *Personalized Learning Material*: Resources to reinforce understanding.
3. *Practice Assignment*: Exercises to solidify learning.

Best regards,
Your School
                    """
                    attachments = [assessment_report_pdf, learning_material_pdf, assignment_pdf]
                    send_email_with_attachments(email_id, subject, body, attachments)
                
                except openai.error.OpenAIError as e:
                    st.error(f"OpenAI API error: {e}")
                except Exception as e:
                    st.error(f"Error in report generation: {e}")
            else:
                st.error("Please provide all required inputs.")

        # Display download buttons for generated reports
        if 'assessment_report_pdf' in st.session_state:
            st.write("Assessment Report")
            with open(st.session_state['assessment_report_pdf'], "rb") as file:
                st.download_button(label="Download Assessment Report as PDF", data=file.read(), file_name=st.session_state['assessment_report_pdf'])

        if 'learning_material_pdf' in st.session_state:
            st.write("Personalized Learning Material")
            with open(st.session_state['learning_material_pdf'], "rb") as file:
                st.download_button(label="Download Learning Material as PDF", data=file.read(), file_name=st.session_state['learning_material_pdf'])

        if 'assignment_pdf' in st.session_state:
            st.write("Personalized Assignment")
            with open(st.session_state['assignment_pdf'], "rb") as file:
                st.download_button(label="Download Assignment as PDF", data=file.read(), file_name=st.session_state['assignment_pdf'])
        
    except Exception as e:
        st.error(f"An error occurred in the Student Assessment Assistant: {e}")











def generate_detailed_answers_with_marking_scheme(question_paper_content):
    """
    Generates detailed answers and a step-by-step marking scheme based on the question paper content.
    """
    prompt = f"""
    Provide detailed answers and a step-by-step marking scheme for the following questions:
    {question_paper_content}
    """
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}]
    )
    return response["choices"][0]["message"]["content"]


def create_answer_sheet_docx(content):
    """Creates a DOCX file for the detailed answer sheet with marking scheme."""
    doc = Document()
    doc.add_heading("Detailed Answer Sheet with Marking Scheme", level=1)
    doc.add_paragraph(content)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def extract_text_from_docx(file):
    doc = Document(file)
    text = [para.text for para in doc.paragraphs if para.text.strip() != ""]
    return "\n".join(text)

# AI function to generate curriculum if not provided
def generate_curriculum(board, subject, class_level):
    prompt = f"""
    You are a curriculum expert. Generate the official curriculum for the {board} board for {subject} in {class_level}.
    Provide a detailed list of topics, subtopics, and learning objectives that align with the official guidelines.
    """
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}]
    )
    return response['choices'][0]['message']['content']

# AI function to check alignment
def check_alignment(assignment_text, curriculum_text):
    prompt = f"""
    You are an expert curriculum alignment checker. Compare the following assignment/lesson plan with the curriculum. 
    Provide a detailed report of alignment, partial alignment, and misalignment with suggestions for improvement.

    Assignment/Lesson Plan Text:
    {assignment_text}

    Curriculum Text:
    {curriculum_text}

    Output the results in the following format:
    - Aligned Topics/Subtopics
    - Partially Aligned Topics/Subtopics
    - Not Aligned Topics/Subtopics
    - Suggestions for Improvement
    """
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}]
    )
    return response['choices'][0]['message']['content']

def generate_detailed_answer_sheet(question_paper_content):
    """
    Generates detailed answers for the provided question paper content.
    """
    try:
        # Check if the question paper content is valid
        if not question_paper_content.strip():
            raise ValueError("Question paper content is empty or invalid.")

        # Prepare the prompt for generating answers
        prompt = f"""
        Provide detailed answers for the following questions:
        {question_paper_content}
        """
        
        # Call the OpenAI API
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}]
        )

        # Extract and return the response content
        if response and "choices" in response and response["choices"]:
            return response["choices"][0]["message"]["content"]
        else:
            raise ValueError("Failed to retrieve a valid response from the AI.")

    except openai.error.OpenAIError as e:
        # Handle OpenAI API errors
        st.error(f"An error occurred with the OpenAI API: {e}")
        return None
    except ValueError as e:
        # Handle validation errors
        st.error(f"Validation Error: {e}")
        return None
    except Exception as e:
        # Handle any other unexpected errors
        st.error(f"An unexpected error occurred: {e}")
        return None

import io

def create_answer_sheet_docx(answer_content):
    """
    Creates a .docx file from the provided answer content.
    Returns the file data as a byte stream.
    """
    doc = Document()
    doc.add_heading("Detailed Answer Sheet", level=1)
    for line in answer_content.split("\n"):
        doc.add_paragraph(line)
    byte_stream = io.BytesIO()
    doc.save(byte_stream)
    byte_stream.seek(0)
    return byte_stream

from fpdf import FPDF

from fpdf import FPDF
import io

from fpdf import FPDF
import io

from fpdf import FPDF
import io

from fpdf import FPDF
import io

from fpdf import FPDF
import io
import textwrap

from fpdf import FPDF

def create_answer_sheet_pdf(answer_content):
    """
    Creates a .pdf file from the provided answer content.
    Returns the file data as a byte stream.
    """
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    pdf.cell(200, 10, txt="Detailed Answer Sheet", ln=True, align='C')
    pdf.ln(10)  # Add a line break

    for line in answer_content.split("\n"):
        pdf.multi_cell(0, 10, line)
    
    byte_stream = io.BytesIO()
    pdf.output(byte_stream)
    byte_stream.seek(0)
    return byte_stream





# Function to generate the sample paper using AI
def generate_sample_paper(board, subject, grade, max_marks, medium):
    prompt = f"""
    You are an expert in educational content creation. Generate a sample paper strictly following the official pattern of the {board} board.
    The subject is '{subject}' for grade '{grade}'. The maximum marks for the paper are {max_marks}, and the medium of the paper is '{medium}'.
    Ensure the following:
    1. The sample paper should strictly follow the board's format and guidelines.
    2. Include all sections and instructions as per the board's official pattern (e.g., MCQs, short answers, long answers).
    3. Adhere to the difficulty level and marks distribution prescribed by the board.
    4. Avoid adding any fields or formats not part of the official pattern.

    Generate content in the {medium} language.
    """
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7,
        max_tokens=4000  #High token limit for a complete sample paper
    )
    
    return response['choices'][0]['message']['content']

# Function to create a Word document with the generated sample paper
def create_docx(content):
    doc = Document()
    doc.add_paragraph(content)
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer




def generate_discussion_prompts(topic, grade, subject, board, medium):
    """
    Generate discussion questions based on the topic, grade, and subject using gpt-3.5-turbo.
    """
    try:
        # Construct the prompt for the GPT model
        prompt = f"""
    Generate five discussion prompts for a classroom lesson on the topic '{topic}', for {grade}, 
    in the subject '{subject}', under the {board} education board. The medium of instruction is {medium}. 
    Ensure the prompts are age-appropriate, engaging, and encourage critical thinking.
    """
        
        # Call the OpenAI GPT model with the required 'messages' parameter
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",  # Use gpt-3.5-turbo
            messages=[{"role": "system", "content": "You are a helpful assistant."},
                      {"role": "user", "content": prompt}],
            max_tokens=1000,
            temperature=0.7,
        )
        # Extract the generated content
        return response["choices"][0]["message"]["content"].strip()
    except Exception as e:
        return f"An error occurred: {e}"

from fpdf import FPDF

def save_prompts_as_pdf(prompts, topic, grade, subject, board, medium):
    """
    Save the generated discussion prompts as a PDF file.
    """
    file_name = f"{topic}_Discussion_Prompts.pdf"
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    pdf.cell(200, 10, txt="Classroom Discussion Prompts", ln=True, align='C')
    pdf.ln(10)  # Add a line break

    pdf.multi_cell(0, 10, f"Topic: {topic}")
    pdf.multi_cell(0, 10, f"Grade: {grade}")
    pdf.multi_cell(0, 10, f"Subject: {subject}")
    pdf.multi_cell(0, 10, f"Board: {board}")
    pdf.multi_cell(0, 10, f"Medium: {medium}")
    pdf.ln(5)  # Add a line break
    pdf.multi_cell(0, 10, f"Prompts:\n{prompts}")

    pdf.output(file_name)
    return file_name

from docx import Document

def save_prompts_as_docx(prompts, topic, grade, subject, board, medium):
    """
    Save the generated discussion prompts as a DOCX file.
    """
    file_name = f"{topic}_Discussion_Prompts.docx"
    document = Document()

    document.add_heading("Classroom Discussion Prompts", level=1)
    document.add_paragraph(f"Topic: {topic}")
    document.add_paragraph(f"Grade: {grade}")
    document.add_paragraph(f"Subject: {subject}")
    document.add_paragraph(f"Board: {board}")
    document.add_paragraph(f"Medium: {medium}")
    document.add_paragraph("\nPrompts:")
    document.add_paragraph(prompts)

    document.save(file_name)
    return file_name

def Question_Paper():
    """
    Generates a question paper with customized number of questions for each question type.
    """
    st.header("Question Paper Creation")
    
    try:
        # Predefined options
        BOARDS = ["CBSE", "ICSE", "State Board", "IB", "Cambridge", "Others"]
        GRADES = [f"Class {i}" for i in range(1, 13)] + ["Others"]
        SUBJECTS = ["Mathematics", "Science", "History", "Geography", "English", "Others"]
        DIFFICULTY_LEVELS = ["Easy", "Medium", "Hard"]
        QUESTION_TYPES = [
            "True/False", "Yes/No", "MCQs", 
            "Short Answers", "Long Answers", 
            "Fill in the Blanks", "Match the Following"
        ]
        MEDIUMS = ["English", "Hindi", "Tamil", "Bengali", "Kiswahili", "Others"]

        # Input fields
        board = st.selectbox("Select Education Board:", BOARDS)
        if board == "Others":
            board = st.text_input("Enter Board Name:")

        standard = st.selectbox("Select Standard/Class:", GRADES)
        if standard == "Others":
            standard = st.text_input("Enter Standard/Class:")

        subject = st.selectbox("Select Subject:", SUBJECTS)
        if subject == "Others":
            subject = st.text_input("Enter Subject Name:")

        medium = st.selectbox("Select Medium:", MEDIUMS)
        if medium == "Others":
            medium = st.text_input("Enter Medium of Instruction:")

        topics = st.text_area("Enter Topics (comma-separated):", placeholder="E.g., Algebra, Photosynthesis")

        total_marks = st.number_input("Enter Total Marks (optional):", min_value=10, step=5)

        time_duration = st.text_input("Enter Time Duration (e.g., 60 minutes, optional):")

        difficulty = st.selectbox("Select Difficulty Level:", DIFFICULTY_LEVELS)

        # Dynamic input for question types and their count
        st.markdown("### Select Question Types and Enter Count for Each")
        question_counts = {}
        for qtype in QUESTION_TYPES:
            count = st.number_input(f"Number of {qtype} Questions:", min_value=0, max_value=50, step=1, key=qtype)
            if count > 0:
                question_counts[qtype] = count

        include_solutions = st.radio("Would you like to include solutions?", ["Yes", "No"]) == "Yes"

        # Generate question paper button
        if st.button("Generate Question Paper"):
            if not board or not standard or not subject or not topics or not question_counts:
                st.warning("Please fill in all required fields and specify at least one question type with a count.")
                return
            
            with st.spinner("Generating question paper..."):
                # Generate content using OpenAI
                content = generate_question_paper(
                    board, standard, subject, topics, total_marks, 
                    time_duration, difficulty, question_counts, 
                    include_solutions
                )
                
                # Display content
                st.write("### Generated Question Paper")
                st.write(content)

                # Save and allow download as DOCX and PDF
                try:
                    docx_file = create_answer_sheet_docx(content)

                    st.download_button(
                        label="Download Question Paper as DOCX",
                        data=docx_file,
                        file_name=f"Question_Paper_{standard}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                except Exception as e:
                    st.error(f"Error saving or downloading files: {e}")

    except Exception as e:
        st.error(f"Error in question paper creation: {e}")


def generate_question_paper(board, standard, subject, topics, total_marks, time_duration, difficulty, question_counts, include_solutions):
    """
    Generates a question paper based on the inputs.
    """
    question_summary = "\n".join([f"{count} {qtype} questions" for qtype, count in question_counts.items()])
    prompt = f"""
    You are an educational content creator. Create a question paper for the {board} board, {standard} class, subject: {subject}.
    Based on the topics: {topics}. The question paper should contain questions as follows:
    {question_summary}. The total marks should be {total_marks}, with a time duration of {time_duration}. 
    The difficulty level should be {difficulty}.
    """
    
    if include_solutions:
        prompt += " Include the solution set."
    else:
        prompt += " Only include the question set without solutions."
    
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "system", "content": prompt}]
        )
        return response['choices'][0]['message']['content']
    except Exception as e:
        raise Exception(f"Error generating question paper: {e}")

# Function to create a PDF from the questions
def create_apdf(content, filename="practice_questions.pdf"):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for line in content.split("\n"):
        pdf.cell(0, 10, line, ln=True)
    pdf.output(filename)
    return filename

# Function to create a DOC file from the questions
def create_adoc(content, filename="practice_questions.docx"):
    doc = Document()
    for line in content.split("\n"):
        doc.add_paragraph(line)
    doc.save(filename)
    return filename


# create Quiz




def create_quiz():
    SUBJECTS = ["Mathematics", "History", "Science", "Geography", "English"]
    DIFFICULTY_LEVELS = ["Beginner", "Intermediate", "Advanced"]
    BOARDS = ["CBSE", "ICSE", "State Board", "IB", "Add New"]

    st.title("AI-Powered Quiz Generator")
    st.write("Generate quizzes dynamically using OpenAI GPT and export them as PDF or DOC!")

    st.header("Quiz Parameters")

    # Board selection with option to add custom boards
    board = st.selectbox("Choose a Board", BOARDS)
    if board == "Add New":
        board = st.text_input("Enter New Board")
        if board and board not in BOARDS:
            BOARDS.append(board)

    # Subject selection with option to add custom subjects
    subject = st.selectbox("Choose a Subject", SUBJECTS + ["Add New"])
    if subject == "Add New":
        subject = st.text_input("Enter New Subject")
        if subject and subject not in SUBJECTS:
            SUBJECTS.append(subject)

    # Topic input
    topic = st.text_input("Enter Topic (e.g., Algebra, World War II)")

    # Difficulty level selection with option to add custom difficulty levels
    difficulty = st.selectbox("Choose Difficulty Level", DIFFICULTY_LEVELS + ["Add New"])
    if difficulty == "Add New":
        difficulty = st.text_input("Enter New Difficulty Level")
        if difficulty and difficulty not in DIFFICULTY_LEVELS:
            DIFFICULTY_LEVELS.append(difficulty)

    # Number of questions
    num_questions = st.slider("Number of Questions", min_value=1, max_value=10, value=5)

    # Include answers option
    include_answers = st.radio("Include Answers in Quiz?", ("Yes", "No")) == "Yes"

    # Generate quiz button
    generate_button = st.button("Generate Quiz")

    if generate_button:
        st.write(f"Generating your quiz for the {board} board... Please wait.")
        quiz_text = generate_quiz(subject, topic, difficulty, num_questions, board, include_answers)
        if quiz_text:
            st.text_area("Generated Quiz", quiz_text, height=300)

            # Buttons to download as PDF or DOC
            pdf_filename = "quiz.pdf"
            doc_filename = "quiz.docx"

            if st.button("Download as PDF"):
                create_pdf(quiz_text, pdf_filename)
                with open(pdf_filename, "rb") as pdf_file:
                    st.download_button(
                        label="Download PDF",
                        data=pdf_file,
                        file_name=pdf_filename,
                        mime="application/pdf"
                    )

            if st.button("Download as DOC"):
                create_doc(quiz_text, doc_filename)
                with open(doc_filename, "rb") as doc_file:
                    st.download_button(
                        label="Download DOC",
                        data=doc_file,
                        file_name=doc_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

def generate_quiz(subject, topic, difficulty, num_questions, board, include_answers):
    """
    Generates quiz content using OpenAI GPT based on the provided parameters.
    """
    prompt = f"""
    Create a quiz for {board} board on the topic "{topic}" in {subject}. 
    Difficulty level: {difficulty}. Number of questions: {num_questions}.
    {"Include answers and explanations." if include_answers else "Do not include answers or explanations."}
    """
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a helpful assistant that generates quizzes."},
                {"role": "user", "content": prompt}
            ]
        )
        return response['choices'][0]['message']['content']
    except Exception as e:
        st.error(f"Error generating quiz: {e}")
        return None


def create_pdf(content, filename):
    """
    Creates a PDF file from the quiz content.
    """
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for line in content.split("\n"):
        pdf.cell(0, 10, line, ln=True)
    pdf.output(filename)


def create_doc(content, filename):
    """
    Creates a DOCX file from the quiz content.
    """
    doc = Document()
    for line in content.split("\n"):
        doc.add_paragraph(line)
    doc.save(filename)



#All functions calling
#Generate Question Paper

def Question_Paper():
    """
    Generates a question paper with customized number of questions for each question type.
    """
    st.header("Question Paper Creation")

    try:
        # Predefined options
        BOARDS = ["CBSE", "ICSE", "State Board", "IB", "Cambridge", "Others"]
        GRADES = [f"Class {i}" for i in range(1, 13)] + ["Others"]
        SUBJECTS = ["Mathematics", "Science", "History", "Geography", "English", "Others"]
        DIFFICULTY_LEVELS = ["Easy", "Medium", "Hard"]
        QUESTION_TYPES = [
            "True/False", "Yes/No", "MCQs",
            "Short Answers", "Long Answers",
            "Fill in the Blanks", "Match the Following"
        ]
        MEDIUMS = ["English", "Hindi", "Tamil", "Bengali", "Kiswahili", "Others"]

        # Input fields
        board = st.selectbox("Select Education Board:", BOARDS)
        if board == "Others":
            board = st.text_input("Enter Board Name:")

        standard = st.selectbox("Select Standard/Class:", GRADES)
        if standard == "Others":
            standard = st.text_input("Enter Standard/Class:")

        subject = st.selectbox("Select Subject:", SUBJECTS)
        if subject == "Others":
            subject = st.text_input("Enter Subject Name:")

        medium = st.selectbox("Select Medium:", MEDIUMS)
        if medium == "Others":
            medium = st.text_input("Enter Medium of Instruction:")

        topics = st.text_area("Enter Topics (comma-separated):", placeholder="E.g., Algebra, Photosynthesis")

        total_marks = st.number_input("Enter Total Marks (optional):", min_value=10, step=5)

        time_duration = st.text_input("Enter Time Duration (e.g., 60 minutes, optional):")

        difficulty = st.selectbox("Select Difficulty Level:", DIFFICULTY_LEVELS)

        # Dynamic input for question types and their count
        st.markdown("### Select Question Types and Enter Count for Each")
        question_counts = {}
        for qtype in QUESTION_TYPES:
            count = st.number_input(f"Number of {qtype} Questions:", min_value=0, max_value=50, step=1, key=qtype)
            if count > 0:
                question_counts[qtype] = count

        include_solutions = st.radio("Would you like to include solutions?", ["Yes", "No"]) == "Yes"

        # Generate question paper button
        if st.button("Generate Question Paper"):
            if not board or not standard or not subject or not topics or not question_counts:
                st.warning("Please fill in all required fields and specify at least one question type with a count.")
                return

            with st.spinner("Generating question paper..."):
                # Generate content using OpenAI
                content = generate_question_paper(
                    board, standard, subject, topics, total_marks,
                    time_duration, difficulty, question_counts,
                    include_solutions
                )

                # Display content
                st.write("### Generated Question Paper")
                st.write(content)

                # Save and allow download as DOCX and PDF
                try:
                    docx_file = create_answer_sheet_docx(content)

                    st.download_button(
                        label="Download Question Paper as DOCX",
                        data=docx_file,
                        file_name=f"Question_Paper_{standard}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                except Exception as e:
                    st.error(f"Error saving or downloading files: {e}")

    except Exception as e:
        st.error(f"Error in question paper creation: {e}")


def generate_question_paper(board, standard, subject, topics, total_marks, time_duration, difficulty, question_counts, include_solutions):
    """
    Generates a question paper based on the inputs.
    """
    question_summary = "\n".join([f"{count} {qtype} questions" for qtype, count in question_counts.items()])
    prompt = f"""
    You are an educational content creator. Create a question paper for the {board} board, {standard} class, subject: {subject}.
    Based on the topics: {topics}. The question paper should contain questions as follows:
    {question_summary}. The total marks should be {total_marks}, with a time duration of {time_duration}. 
    The difficulty level should be {difficulty}.
    """

    if include_solutions:
        prompt += " Include the solution set."
    else:
        prompt += " Only include the question set without solutions."

    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "system", "content": prompt}]
        )
        return response['choices'][0]['message']['content']
    except Exception as e:
        raise Exception(f"Error generating question paper: {e}")


def create_apdf(content, filename="practice_questions.pdf"):
    """
    Function to create a PDF from the questions.
    """
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for line in content.split("\n"):
        pdf.cell(0, 10, line, ln=True)
    pdf.output(filename)
    return filename


def create_adoc(content, filename="practice_questions.docx"):
    """
    Function to create a DOC file from the questions.
    """
    doc = Document()
    for line in content.split("\n"):
        doc.add_paragraph(line)
    doc.save(filename)
    return filename

#sample paper generator function
def Sample_Paper():
            st.title("Board-Specific Sample Paper Generator")
            
            # Predefined lists for selection
            predefined_boards = ["CBSE", "ICSE", "IB", "Cambridge", "State Board"]
            predefined_subjects = ["Mathematics", "Science", "English", "History"]
            predefined_grades = ["Grade 1", "Grade 5", "Grade 10", "Grade 12"]
            predefined_mediums = ["English", "Hindi", "French"]

            # Board selection
            board = st.selectbox("Select Board:", predefined_boards + ["Other"])
            if board == "Other":
                board = st.text_input("Enter Board:")

            # Subject selection
            subject = st.selectbox("Select Subject:", predefined_subjects + ["Other"])
            if subject == "Other":
                subject = st.text_input("Enter Subject:")

            # Grade selection
            grade = st.selectbox("Select Grade:", predefined_grades + ["Other"])
            if grade == "Other":
                grade = st.text_input("Enter Grade:")

            # Medium selection
            medium = st.selectbox("Select Medium:", predefined_mediums + ["Other"])
            if medium == "Other":
                medium = st.text_input("Enter Medium:")

            # Maximum marks input
            max_marks = st.number_input("Maximum Marks:", min_value=10, value=100)

            # Generate sample paper
            if st.button("Generate Sample Paper"):
                if not board or not subject or not grade or not medium:
                    st.error("Please fill in all the required fields.")
                else:
                    try:
                        # Generate sample paper content
                        sample_paper_content = generate_sample_paper(board, subject, grade, max_marks, medium)

                        # Display the sample paper content on the screen
                        st.subheader("Generated Sample Paper:")
                        st.write(sample_paper_content)

                        # Create downloadable DOCX and PDF files using reusable functions
                        docx_file = create_answer_sheet_docx(sample_paper_content)
                        
                        # Provide download buttons
                        st.download_button(
                            label="Download Sample Paper (DOCX)",
                            data=docx_file,
                            file_name=f"{board}_{subject}_{grade}_Sample_Paper.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

                        
                    except Exception as e:
                        st.error(f"An error occurred: {e}")
#generate assignment
def Assignment():
    st.title("Assignment Generator")

    # Predefined options
    boards = ["CBSE", "ICSE", "IB", "State Board", "Others"]
    subjects = ["Mathematics", "Science", "English", "Others"]
    grades = ["Grade 1", "Grade 5", "Grade 10", "Grade 12", "Others"]

    # Select inputs
    board = st.selectbox("Select Board", boards)
    if board == "Others":
        board = st.text_input("Enter Board Name:")

    subject = st.selectbox("Select Subject", subjects)
    if subject == "Others":
        subject = st.text_input("Enter Subject Name:")

    grade = st.selectbox("Select Grade", grades)
    if grade == "Others":
        grade = st.text_input("Enter Grade/Level:")

    # Topics input: Allow adding multiple topics
    topics = st.text_area("Enter Topics (comma-separated):", placeholder="e.g., Algebra, Photosynthesis")

    # Optional marks input
    marks = st.number_input("Enter Maximum Marks (Optional):", min_value=10, max_value=100, step=5)

    # Generate Assignment button
    if st.button("Generate Assignment"):
        # Validation for required fields
        if not board or not subject or not grade or not topics.strip():
            st.warning("Please fill in all required fields (Board, Subject, Grade, and Topics).")
        else:
            # Prepare assignment prompt
            assignment_prompt = f"Generate an assignment for {subject} for {grade}."
            assignment_prompt += f" Follow the {board} curriculum."
            assignment_prompt += f" Focus on the topics: {topics}."
            if marks:
                assignment_prompt += f" The assignment should be worth {marks} marks."

            # Call OpenAI API
            with st.spinner("Generating assignment..."):
                try:
                    assignment = call_openai_api(assignment_prompt, max_tokens=1500)
                    st.write("### Generated Assignment:")
                    st.write(assignment)
                except Exception as e:
                    st.error(f"An error occurred: {e}")



#create quiz function
def Quiz():
    SUBJECTS = ["Mathematics", "History", "Science", "Geography", "English"]
    DIFFICULTY_LEVELS = ["Beginner", "Intermediate", "Advanced"]
    BOARDS = ["CBSE", "ICSE", "State Board", "IB", "Add New"]

    st.title("AI-Powered Quiz Generator")
    st.write("Generate quizzes dynamically using OpenAI GPT and export them as PDF or DOC!")

    st.header("Quiz Parameters")

    # Board selection with option to add custom boards
    board = st.selectbox("Choose a Board", BOARDS)
    if board == "Add New":
        board = st.text_input("Enter New Board")
        if board and board not in BOARDS:
            BOARDS.append(board)

    # Subject selection with option to add custom subjects
    subject = st.selectbox("Choose a Subject", SUBJECTS + ["Add New"])
    if subject == "Add New":
        subject = st.text_input("Enter New Subject")
        if subject and subject not in SUBJECTS:
            SUBJECTS.append(subject)

    # Topic input
    topic = st.text_input("Enter Topic (e.g., Algebra, World War II)")

    # Difficulty level selection with option to add custom difficulty levels
    difficulty = st.selectbox("Choose Difficulty Level", DIFFICULTY_LEVELS + ["Add New"])
    if difficulty == "Add New":
        difficulty = st.text_input("Enter New Difficulty Level")
        if difficulty and difficulty not in DIFFICULTY_LEVELS:
            DIFFICULTY_LEVELS.append(difficulty)

    # Number of questions
    num_questions = st.slider("Number of Questions", min_value=1, max_value=10, value=5)

    # Include answers option
    include_answers = st.radio("Include Answers in Quiz?", ("Yes", "No")) == "Yes"

    # Generate quiz button
    generate_button = st.button("Generate Quiz")

    if generate_button:
        st.write(f"Generating your quiz for the {board} board... Please wait.")
        quiz_text = generate_quiz(subject, topic, difficulty, num_questions, board, include_answers)
        if quiz_text:
            st.text_area("Generated Quiz", quiz_text, height=300)

            # Buttons to download as PDF or DOC
            pdf_filename = "quiz.pdf"
            doc_filename = "quiz.docx"

            if st.button("Download as PDF"):
                create_pdf(quiz_text, pdf_filename)
                with open(pdf_filename, "rb") as pdf_file:
                    st.download_button(
                        label="Download PDF",
                        data=pdf_file,
                        file_name=pdf_filename,
                        mime="application/pdf"
                    )

            if st.button("Download as DOC"):
                create_doc(quiz_text, doc_filename)
                with open(doc_filename, "rb") as doc_file:
                    st.download_button(
                        label="Download DOC",
                        data=doc_file,
                        file_name=doc_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )


def generate_quiz(subject, topic, difficulty, num_questions, board, include_answers):
    """
    Generates quiz content using OpenAI GPT based on the provided parameters.
    """
    prompt = f"""
    Create a quiz for {board} board on the topic "{topic}" in {subject}. 
    Difficulty level: {difficulty}. Number of questions: {num_questions}.
    {"Include answers and explanations." if include_answers else "Do not include answers or explanations."}
    """
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a helpful assistant that generates quizzes."},
                {"role": "user", "content": prompt}
            ]
        )
        return response['choices'][0]['message']['content']
    except Exception as e:
        st.error(f"Error generating quiz: {e}")
        return None


def create_pdf(content, filename):
    """
    Creates a PDF file from the quiz content.
    """
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for line in content.split("\n"):
        pdf.cell(0, 10, line, ln=True)
    pdf.output(filename)


def create_doc(content, filename):
    """
    Creates a DOCX file from the quiz content.
    """
    doc = Document()
    for line in content.split("\n"):
        doc.add_paragraph(line)
    doc.save(filename)





#Generate Lesson Plan
def Lesson_Plans():
    create_lesson_plan()


# Generate Image Based Questions
def Image_Based_Questions():
    generate_image_based_questions()

#Generate Paragraph based Questions
#def Paragraph_Based_Question():

#Generate Classroom Discussion Prompt
def Classroom_Discussion_Prompter():
    st.title("Classroom Discussion Prompter")
    st.markdown("### Make your lessons more interactive with creative discussion prompts!")

    try:
        # Predefined options for boards, mediums, and subjects
        predefined_boards = ["CBSE", "ICSE", "State Board", "IB", "Cambridge", "Others"]
        predefined_mediums = ["English", "Kiswahili", "Hindi", "Gujarati", "French", "Others"]
        predefined_subjects = ["Mathematics", "Science", "History", "Geography", "English", "Others"]

        # User inputs
        topic = st.text_input("Enter the lesson topic:", placeholder="e.g., Photosynthesis, World War II, Algebra")
        grade = st.selectbox(
            "Select the grade level:",
            ["Grade 1", "Grade 2", "Grade 3", "Grade 4", "Grade 5", "Grade 6", "Grade 7",
             "Grade 8", "Grade 9", "Grade 10", "Grade 11", "Grade 12"]
        )

        # Dropdown for subject with "Add Another" functionality
        subject = st.selectbox("Select the subject:", predefined_subjects + ["Add Another"])
        if subject == "Add Another":
            subject = st.text_input("Enter your subject:")

        # Dropdown fields for Board and Medium
        board = st.selectbox("Select the education board:", predefined_boards)
        if board == "Others":
            board = st.text_input("Enter your board name:")

        medium = st.selectbox("Select the medium of instruction:", predefined_mediums)
        if medium == "Others":
            medium = st.text_input("Enter your preferred medium:")

        # Generate discussion prompts
        if st.button("Generate Discussion Prompts"):
            if topic and grade and subject and board and medium:
                with st.spinner("Generating discussion prompts..."):
                    try:
                        # Call a function to generate discussion prompts
                        prompts = generate_discussion_prompts(topic, grade, subject, board, medium)
                        
                        if prompts.strip():
                            st.subheader("Suggested Discussion Prompts:")
                            st.write(prompts)

                            # Save prompts as DOCX and provide a download option
                            docx_file = create_answer_sheet_docx(prompts)
                            st.download_button(
                                label="Download as DOCX",
                                data=docx_file,
                                file_name=f"{topic}_Discussion_Prompts.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        else:
                            st.error("Failed to generate valid discussion prompts.")
                    except Exception as e:
                        st.error(f"An error occurred during prompt generation: {e}")
            else:
                st.error("Please fill in all the fields to generate discussion prompts.")
    except Exception as e:
        st.error(f"An unexpected error occurred: {e}")



# Module2
 # Add logic for assessment assistance
def Assessment_Assistant():
    # Add logic for answer sheet generation   
    student_assessment_assistant()



def Generate_Answer_Sheets():
    # Add logic for marking scheme generation
    st.title("Answer Sheet Generator")
    st.markdown("""
        1. Upload your question paper in .docx format.
        2. The AI will generate **detailed answers** for each question.
        3. View the generated answer sheet on screen.
        4. Download the generated answer sheet as a .docx or .pdf file.
    """)

    # File uploader for the question paper
    uploaded_file = st.file_uploader("Upload Question Paper (.docx)", type=["docx"])

    if uploaded_file and st.button("Generate Detailed Answer Sheet"):
        try:
            # Read the content of the uploaded question paper
            question_paper_content = read_docx(uploaded_file)

            if not question_paper_content.strip():
                st.error("The uploaded question paper appears to be empty.")
                return

            # Generate detailed answer sheet content
            answer_sheet_content = generate_detailed_answer_sheet(question_paper_content)

            if answer_sheet_content:
                # Display the answer sheet content on the screen
                st.subheader("Generated Answer Sheet:")
                st.write(answer_sheet_content)

                # Create downloadable .docx file
                docx_file = create_answer_sheet_docx(answer_sheet_content)
                st.download_button(
                    "Download Detailed Answer Sheet (DOCX)",
                    data=docx_file,
                    file_name="Detailed_Answer_Sheet.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            else:
                st.error("Failed to generate the detailed answer sheet.")
        except Exception as e:
            st.error(f"An error occurred: {e}")


# Placeholder for read_docx function
def read_docx(file):
    """
    Reads the content of a .docx file and returns it as a string.
    """
    try:
        doc = Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        raise Exception(f"Error reading DOCX file: {e}")


    # Placeholder for generate_detailed_answer_sheet function
    def generate_detailed_answer_sheet(question_paper_content):
        """
        Generates a detailed answer sheet based on the question paper content.
        """
        prompt = f"Provide detailed answers for the following questions:\n{question_paper_content}"
        try:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant for generating answer sheets."},
                    {"role": "user", "content": prompt}
                ]
            )
            return response['choices'][0]['message']['content']
        except Exception as e:
            raise Exception(f"Error generating detailed answer sheet: {e}")


# Placeholder for create_answer_sheet_docx function
def create_answer_sheet_docx(content):
    """
    Creates a DOCX file from the provided content.
    """
    try:
        doc = Document()
        for line in content.split("\n"):
            doc.add_paragraph(line)
        doc_file = io.BytesIO()
        doc.save(doc_file)
        doc_file.seek(0)
        return doc_file
    except Exception as e:
        raise Exception(f"Error creating DOCX file: {e}")


def Marking_Scheme():
    # Add logic for analyzing reports
    st.title("Marking Scheme Generator")
    st.markdown("""
        1. Upload your question paper in .docx format.
        2. The AI will generate **detailed answers** and a **step-by-step marking scheme**.
        3. View the content on the screen and download it as a **DOCX** or **PDF** file.
    """)

    # File uploader for question paper
    uploaded_file = st.file_uploader("Upload Question Paper (.docx)", type=["docx"])

    if uploaded_file and st.button("Generate Answer Sheet & Marking Scheme"):
        with st.spinner("Generating detailed answers and marking scheme..."):
            try:
                # Read the uploaded question paper
                question_paper_content = read_docx(uploaded_file)

                if not question_paper_content.strip():
                    st.error("The uploaded question paper appears to be empty.")
                    return

                # Generate detailed answers and marking scheme
                detailed_content = generate_detailed_answers_with_marking_scheme(question_paper_content)

                if detailed_content:
                    # Display the content on the screen
                    st.subheader("Generated Answer Sheet and Marking Scheme:")
                    st.write(detailed_content)

                    # Create downloadable DOCX file
                    docx_file = create_answer_sheet_docx(detailed_content)
                    st.download_button(
                        label="Download Answer Sheet with Marking Scheme (DOCX)",
                        data=docx_file,
                        file_name="Detailed_Answer_Sheet_with_Marking_Scheme.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.error("Failed to generate the detailed answers and marking scheme.")
            except Exception as e:
                st.error(f"An error occurred: {e}")


# Placeholder for read_docx function
def read_docx(file):
    """
    Reads the content of a .docx file and returns it as a string.
    """
    try:
        doc = Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        raise Exception(f"Error reading DOCX file: {e}")


# Placeholder for generate_detailed_answers_with_marking_scheme function
def generate_detailed_answers_with_marking_scheme(question_paper_content):
    """
    Generates detailed answers and a marking scheme based on the question paper content.
    """
    prompt = f"""
    Generate detailed answers and a marking scheme for the following question paper:
    {question_paper_content}
    """
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a helpful assistant for generating marking schemes."},
                {"role": "user", "content": prompt}
            ]
        )
        return response['choices'][0]['message']['content']
    except Exception as e:
        raise Exception(f"Error generating marking scheme: {e}")


# Placeholder for create_answer_sheet_docx function
def create_answer_sheet_docx(content):
    """
    Creates a DOCX file from the provided content.
    """
    try:
        doc = Document()
        for line in content.split("\n"):
            doc.add_paragraph(line)
        doc_file = io.BytesIO()
        doc.save(doc_file)
        doc_file.seek(0)
        return doc_file
    except Exception as e:
        raise Exception(f"Error creating DOCX file: {e}")


def Analyze_Reports():
    report_analysis_with_openai()


#Module3

# Add logic for curriculum generation
def Curriculum_Generator():
    st.title("Curriculum Generator")

    # Predefined options for Board, Grade, and Subject
    predefined_boards = ["CBSE", "ICSE", "IB", "State Board", "Others"]
    predefined_grades = ["Grade 1", "Grade 5", "Grade 10", "Grade 12", "Others"]
    predefined_subjects = ["Mathematics", "Science", "History", "Others"]

    # Board selection with "Others" handling
    board = st.selectbox("Select Board", predefined_boards)
    if board == "Others":
        board = st.text_input("Enter Board Name:")

    # Grade selection with "Others" handling
    grade = st.selectbox("Select Grade", predefined_grades)
    if grade == "Others":
        grade = st.text_input("Enter Grade/Level:")

    # Subject selection with "Others" handling
    subject = st.selectbox("Select Subject", predefined_subjects)
    if subject == "Others":
        subject = st.text_input("Enter Subject Name:")

    # Generate Curriculum
    if st.button("Generate Curriculum"):
        if board.strip() and grade.strip() and subject.strip():
            with st.spinner("Generating curriculum..."):
                try:
                    # Call OpenAI API to generate the curriculum
                    curriculum = call_openai_api(
                        f"Generate a detailed curriculum for {subject} for {board}, {grade}.",
                        max_tokens=1500,
                    )

                    if curriculum.strip():
                        # Display the curriculum on the screen
                        st.subheader("Generated Curriculum")
                        for line in curriculum.split("\n"):  # Ensure long lines are broken
                            st.write(line)

                        # Create downloadable DOCX file
                        docx_file = create_answer_sheet_docx(curriculum)

                        # Provide download button
                        st.download_button(
                            label="Download Curriculum (DOCX)",
                            data=docx_file,
                            file_name=f"{board}_{grade}_{subject}_Curriculum.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    else:
                        st.error("Failed to generate a valid curriculum. Please try again.")
                except Exception as e:
                    st.error(f"An error occurred: {e}")
        else:
            st.error("Please fill in all the required fields.")

def Alignment_Checker():
    st.title("Curriculum Alignment Checker")

    # User inputs for Board, Subject, and Class/Grade
    board = st.selectbox("Select Board", ["CBSE", "ICSE", "IGCSE", "State Board", "Enter Manually"])
    if board == "Enter Manually":
        board = st.text_input("Enter Board Name")

    subject = st.selectbox("Select Subject", ["Mathematics", "Science", "English", "Social Studies", "Enter Manually"])
    if subject == "Enter Manually":
        subject = st.text_input("Enter Subject Name")

    class_level = st.selectbox("Select Class/Grade", [f"Class {i}" for i in range(1, 13)] + ["Enter Manually"])
    if class_level == "Enter Manually":
        class_level = st.text_input("Enter Class/Grade")

    # File upload inputs
    st.write("Upload the Assignment, Lesson Plan, or Question Paper:")
    assignment_file = st.file_uploader("Upload DOCX file", type=["docx"])

    st.write("Upload the Curriculum File (Optional):")
    curriculum_file = st.file_uploader("Upload Curriculum DOCX file", type=["docx"])

    # Generate Alignment Report
    if st.button("Check Curriculum Alignment"):
        try:
            # Extract text from uploaded files
            assignment_text = extract_text_from_docx(assignment_file) if assignment_file else ""
            curriculum_text = extract_text_from_docx(curriculum_file) if curriculum_file else ""

            # Generate curriculum if not provided
            if not curriculum_text and board.strip() and subject.strip() and class_level.strip():
                curriculum_text = generate_curriculum(board, subject, class_level)

            # Check alignment
            if assignment_text and curriculum_text:
                alignment_report = check_alignment(assignment_text, curriculum_text)

                # Display alignment report on the screen
                st.subheader("Alignment Report")
                st.write(alignment_report)

                # Create downloadable DOCX file
                docx_file = create_answer_sheet_docx(alignment_report)

                # Provide download button
                st.download_button(
                    label="Download Alignment Report (DOCX)",
                    data=docx_file,
                    file_name="alignment_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.error("Please upload the necessary files or fill in all required fields.")
        except Exception as e:
            st.error(f"An error occurred: {e}")

def call_openai_api(prompt, max_tokens=1500):
    """
    Calls OpenAI API with the given prompt and returns the generated text.
    """
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "system", "content": prompt}]
        )
        return response['choices'][0]['message']['content']
    except Exception as e:
        raise Exception(f"Error calling OpenAI API: {e}")


def extract_text_from_docx(file):
    """
    Extracts text from a .docx file.
    """
    try:
        doc = Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        raise Exception(f"Error extracting text from DOCX file: {e}")


def generate_curriculum(board, subject, class_level):
    """
    Generates a default curriculum based on board, subject, and class level.
    """
    prompt = f"Create a curriculum for {board} board in {subject} for {class_level}."
    return call_openai_api(prompt)


def check_alignment(assignment_text, curriculum_text):
    """
    Checks alignment between the assignment and curriculum text.
    """
    prompt = f"""
    Compare the following assignment with the curriculum and identify alignment or gaps:
    Assignment: {assignment_text}
    Curriculum: {curriculum_text}
    """
    return call_openai_api(prompt)


def create_answer_sheet_docx(content):
    """
    Creates a DOCX file from the provided content.
    """
    try:
        doc = Document()
        for line in content.split("\n"):
            doc.add_paragraph(line)
        doc_file = io.BytesIO()
        doc.save(doc_file)
        doc_file.seek(0)
        return doc_file
    except Exception as e:
        raise Exception(f"Error creating DOCX file: {e}")


    
#Module4
 # Add logic for text generation
def Text_Generation():
    st.title("AI Writing Assistant")

    # Input fields for prompt, tone, and length
    prompt = st.text_area("Enter your prompt for text generation:")
    tone = st.selectbox("Select Tone", ["Formal", "Casual", "Creative", "Professional"])
    length = st.slider("Select Content Length", 50, 1000, step=50)

    # Generate text button
    if st.button("Generate Text"):
        if prompt.strip():
            try:
                # Call OpenAI API to generate text
                generated_text = call_openai_api(
                    f"Tone: {tone}\nLength: {length} words\nPrompt: {prompt}"
                )
                if generated_text.strip():
                    st.subheader("Generated Text:")
                    st.write(generated_text)
                else:
                    st.error("No text was generated. Please try again with a different prompt.")
            except Exception as e:
                st.error(f"An error occurred: {e}")
        else:
            st.error("Please enter a prompt for text generation.")
def Text_Generation():
    st.title("AI Writing Assistant")

    # Input fields for prompt, tone, and length
    prompt = st.text_area("Enter your prompt for text generation:")
    tone = st.selectbox("Select Tone", ["Formal", "Casual", "Creative", "Professional"])
    length = st.slider("Select Content Length", 50, 1000, step=50)

    # Generate text button
    if st.button("Generate Text"):
        if prompt.strip():
            try:
                # Call OpenAI API to generate text
                generated_text = call_openai_api(
                    f"Tone: {tone}\nLength: {length} words\nPrompt: {prompt}"
                )
                if generated_text.strip():
                    st.subheader("Generated Text:")
                    st.write(generated_text)
                else:
                    st.error("No text was generated. Please try again with a different prompt.")
            except Exception as e:
                st.error(f"An error occurred: {e}")
        else:
            st.error("Please enter a prompt for text generation.")

def Advanced_Editing():
    st.title("Advanced Text Editing")

    # Input fields for text and editing type
    text = st.text_area("Enter Text for Editing:")
    option = st.selectbox("Choose Editing Type", ["Grammar Check", "Rewrite", "Summarize"])

    # Edit text button
    if st.button("Edit Text"):
        if text.strip():
            try:
                # Call OpenAI API for editing
                edited_text = call_openai_api(
                    f"Perform {option} on the following text:\n{text}", max_tokens=500
                )

                if edited_text.strip():
                    st.subheader("Edited Text:")
                    st.write(edited_text)

                    # Create downloadable DOCX file
                    try:
                        docx_file = create_answer_sheet_docx(edited_text)
                        st.download_button(
                            label="Download Edited Text (DOCX)",
                            data=docx_file,
                            file_name="Edited_Text.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    except Exception as e:
                        st.error(f"Error creating downloadable files: {e}")
                else:
                    st.error("The API returned an empty response. Please try again.")
            except Exception as e:
                st.error(f"An unexpected error occurred: {e}")



def main_app():
    """
    Main application function that controls the display and functionality of each app section.
    Provides error handling and UI feedback for each module.
    """
    if st.button("Logout"):
        logout()

    try:
        client_config = st.session_state.get('client_config')

        # Display client logo and name with theme color and style
        if client_config:
            st.markdown(f"""
                <div style="text-align: center; background: linear-gradient(180deg, #6A5ACD, {client_config['theme_color']}); padding: 10px 0; border-radius: 8px;">
                    <h2 style="margin: 0; font-size: 24px; color: white;">{client_config['name']}</h2>
                </div>
            """, unsafe_allow_html=True)

        # Sidebar with improved styling for module selection
        st.sidebar.markdown("""
            <style>
                .sidebar-title { font-size: 22px; color: #4B0082; margin-bottom: 15px; text-align: center; }
            </style>
        """, unsafe_allow_html=True)

       # st.title("Educational Tools Hub")
       # st.write("Welcome to the AI-powered educational tools platform! Use the sidebar to navigate through the modules.")

        # Sidebar task selection
        task = st.sidebar.selectbox(
            "Select a Module",
            [
                "Home",
                "Educational Content Creation",
                "Student Assessment & Evaluation",
                "Curriculum & Alignment",
                "Advanced Editing & Text Generation",
            ]
        )
         
        # Function to call OpenAI API based on selected model
        def call_openai_api(prompt, max_tokens=1000):
            response = openai.ChatCompletion.create(
                model=ai_model,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=max_tokens,
            )
            return response["choices"][0]["message"]["content"]

        # Content based on selected module
        if task == "Home":
            show_home()
            
       
        # Educational Content Creation
        elif task == "Educational Content Creation":
            subtask = st.sidebar.radio("Select a Submodule", [
                "Generate Question Paper",
                "Generate Sample Paper",
                "Generate Assignment",
                "Generate Quiz",
                "Generate Lesson Plans",
                "Generate Image-Based Questions",
                "Generate Paragraph Based Question",
                "Generate Classroom Discussion Prompter"
            ])
            if subtask == "Generate Question Paper":
                #st.header("Question Paper Generator")
                Question_Paper()

            elif subtask == "Generate Sample Paper":
               # st.header("Sample Paper Generator")
                Sample_Paper()

            elif subtask == "Generate Assignment":
                #st.header("Assignment Generator")
                Assignment()

            elif subtask == "Generate Quiz":
                #st.header("Quiz Generator")
                Quiz()

            elif subtask == "Generate Lesson Plans":
                #st.header("Lesson Plan Generator")
                Lesson_Plans()

            elif subtask == "Generate Image-Based Questions":
                #st.header("Image-Based Question Generator")
                Image_Based_Questions()

            elif subtask == "Generate Paragraph Based Question":
                #st.header("Paragraph-Based Question Generator")
                Paragraph_Based_Questions()

            elif subtask == "Generate Classroom Discussion Prompter":
                #st.header("Classroom Discussion Prompter")
                Classroom_Discussion_Prompter()

        # Student Assessment & Evaluation
        elif task == "Student Assessment & Evaluation":
            subtask = st.sidebar.radio("Select a Submodule", [
                "Student Assessment Assistant",
                "Generate Answer Sheets",
                "Marking Scheme Generator",
                "Analyze Reports",
                "Grading",
                "Performance Graph"
            ])
            if subtask == "Student Assessment Assistant":
                #st.header("Student Assessment Assistant")
                Assessment_Assistant()

            elif subtask == "Generate Answer Sheets":
                #st.header("Answer Sheet Generator")
                Generate_Answer_Sheets()

            elif subtask == "Marking Scheme Generator":
                #st.header("Marking Scheme Generator")
                Marking_Scheme()

            elif subtask == "Analyze Reports":
                #st.header("Assessment Report Analyzer")
                Analyze_Reports()

            elif subtask == "Grading":
                grading()

            elif subtask == "Performance Graph":
                performance_graph()

        # Curriculum & Alignment
        elif task == "Curriculum & Alignment":
            subtask = st.sidebar.radio("Select a Submodule", [
                "Curriculum Generator",
                "Alignment Checker"
            ])
            if subtask == "Curriculum Generator":
                #st.header("Curriculum Generator")
                Curriculum_Generator()

            elif subtask == "Alignment Checker":
                #st.header("Curriculum Alignment Checker")
                Alignment_Checker()

        # Advanced Editing & Text Generation
        elif task == "Advanced Editing & Text Generation":

            subtask = st.sidebar.radio("Select a Submodule", [
                "Text Generation",
                "Advanced Editing"
            ])
            if subtask == "Text Generation":
                #st.header("Text Generator")
                Text_Generation()

            elif subtask == "Advanced Editing":
                #st.header("Advanced Editing Tools")
                Advanced_Editing()

    # Dynamic Theme Selection
        theme = st.sidebar.selectbox(
            "Choose Theme", ["Default", "Dark", "Light", "Custom"]
        )

        # Apply the selected theme dynamically
        if theme == "Default":
            st.markdown("""
                <style>
                    body { background-color: #f0f2f6; color: #000000; }
                    .stApp { background-color: #f0f2f6; }
                </style>
            """, unsafe_allow_html=True)
        elif theme == "Dark":
            st.markdown("""
                <style>
                    body { background-color: #121212; color: #e0e0e0; }
                    .stApp { background-color: #121212; }
                    h1, h2, h3, h4, h5, h6 { color: #ffffff; }
                </style>
            """, unsafe_allow_html=True)
        elif theme == "Light":
            st.markdown("""
                <style>
                    body { background-color: #ffffff; color: #000000; }
                    .stApp { background-color: #ffffff; }
                    h1, h2, h3, h4, h5, h6 { color: #000000; }
                </style>
            """, unsafe_allow_html=True)
        elif theme == "Custom":
            bg_color = st.sidebar.color_picker("Select Background Color", "#ffffff")
            text_color = st.sidebar.color_picker("Select Text Color", "#000000")
            st.markdown(f"""
                <style>
                    body {{ background-color: {bg_color}; color: {text_color}; }}
                    .stApp {{ background-color: {bg_color}; }}
                    h1, h2, h3, h4, h5, h6 {{ color: {text_color}; }}
                </style>
            """, unsafe_allow_html=True)

        # AI Model Selection
        ai_model = st.sidebar.radio("Choose AI Model", ["gpt-3.5-turbo", "gpt-4"])
    
    
    except KeyError as e:
        st.error(f"Configuration error: {e}. Please log in again.")
    except Exception as e:
        st.error(f"An unexpected error occurred in the main app: {e}")





    




import streamlit as st

def landing_page():
    # Set page configuration for better display
    st.set_page_config(page_title="EduPro - Transform Your Teaching Experience", page_icon="📘", layout="wide")

    # Apply custom CSS for the landing page
    st.markdown(
        """
        <style>
            /* Centered main container styling */
            .container {
                text-align: center;
                margin-top: 5%; /* Reduced top margin for better alignment */
                max-width: 800px;
                margin-left: auto;
                margin-right: auto;
            }

            /* Logo styling */
            .logo {
                font-size: 36px;
                font-weight: 700;
                color: #5A5DF5;
                display: inline-flex;
                align-items: center;
                gap: 10px;
                font-family: 'Arial', sans-serif;
                margin-bottom: 20px;
            }
            .logo-icon {
                font-size: 40px;
                color: #9933FF;
            }

            /* Headline styling */
            .headline {
                font-size: 44px;
                font-weight: 800;
                color: #4B0082;
                line-height: 1.2;
                margin-bottom: 20px;
                font-family: 'Arial', sans-serif;
            }
            .highlight {
                color: #5A5DF5;
                font-weight: 800;
            }

            /* Subtext styling */
            .subtext {
                font-size: 20px;
                color: #666;
                margin-bottom: 40px;
                font-family: 'Arial', sans-serif;
            }

            /* Button container */
            .button-container {
                display: flex;
                justify-content: center;
                margin-top: 30px;
            }
            /* Button styling */
            .stButton>button {
                font-size: 18px;
                font-weight: 600;
                padding: 15px 50px;
                border-radius: 10px;
                transition: all 0.3s ease;
                cursor: pointer;
                border: none;
                font-family: 'Arial', sans-serif;
            }
            /* Primary button style */
            .stButton>button {
                background: linear-gradient(90deg, #5A5DF5, #9933FF);
                color: white;
                font-weight: 700;
                box-shadow: 0px 6px 15px rgba(90, 93, 245, 0.4);
            }
            .stButton>button:hover {
                background: linear-gradient(90deg, #4A4CE3, #802DD4);
                box-shadow: 0px 8px 18px rgba(90, 93, 245, 0.6);
            }

            /* Category cards */
            .category-cards {
                display: flex;
                justify-content: center;
                gap: 30px;
                margin-top: 40px;
                flex-wrap: wrap;
            }
            .category-card {
                background-color: #F8F9FC;
                border-radius: 12px;
                padding: 20px;
                box-shadow: 0px 4px 10px rgba(0, 0, 0, 0.1);
                width: 250px;
                text-align: center;
                transition: transform 0.3s ease, box-shadow 0.3s ease;
            }
            .category-card:hover {
                transform: translateY(-10px);
                box-shadow: 0px 6px 15px rgba(0, 0, 0, 0.2);
            }
            .category-icon {
                font-size: 40px;
                color: #5A5DF5;
                margin-bottom: 10px;
            }
            .category-title {
                font-size: 18px;
                font-weight: 700;
                margin-bottom: 5px;
                color: #333;
            }
            .category-description {
                font-size: 14px;
                color: #666;
            }
        </style>
        """,
        unsafe_allow_html=True,
    )

    # Landing page content
    st.markdown(
        """
        <div class="container">
            <div class="logo">
                <span class="logo-icon">📘</span> EduPro
            </div>
            <div class="headline">
                Transform Your <span class="highlight">Teaching Experience</span>
            </div>
            <div class="subtext">
                Empowering educators with cutting-edge tools to create engaging content, evaluate students effectively, and streamline classroom activities.
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # Button container
    st.markdown('<div class="button-container">', unsafe_allow_html=True)
    if st.button("Get started ➔", key="get_started", help="Click to proceed to the login page"):
        st.session_state['page'] = 'login'  # Set page to 'login' to navigate to login page
    st.markdown('</div>', unsafe_allow_html=True)

    # Categories Section
    st.markdown(
        """
        <div class="category-cards">
            <div class="category-card">
                <div class="category-icon">📚</div>
                <div class="category-title">Content Creation</div>
                <div class="category-description">Design quizzes, lessons, and materials with ease.</div>
            </div>
            <div class="category-card">
                <div class="category-icon">📝</div>
                <div class="category-title">Student Assessments</div>
                <div class="category-description">Evaluate student performance with precision.</div>
            </div>
            <div class="category-card">
                <div class="category-icon">📊</div>
                <div class="category-title">Report Analysis</div>
                <div class="category-description">Generate insightful reports and visualizations.</div>
            </div>
            <div class="category-card">
                <div class="category-icon">💡</div>
                <div class="category-title">AI Assistance</div>
                <div class="category-description">Leverage AI to make informed decisions effortlessly.</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

# Ensure session state for 'page' is initialized
if 'page' not in st.session_state:
    st.session_state['page'] = 'home'  # Set a default value, e.g., 'home'




import pandas as pd
import streamlit as st

# Grading logic based on percentage
def calculate_grade_grading(percentage):
    if percentage >= 90:
        return "A+"
    elif percentage >= 80:
        return "A"
    elif percentage >= 70:
        return "B"
    elif percentage >= 60:
        return "C"
    elif percentage >= 50:
        return "D"
    else:
        return "F"

# Comment logic based on grade
def get_comment_grading(grade):
    comments = {
        "A+": "Excellent performance!",
        "A": "Very good! Keep it up.",
        "B": "Good effort, aim higher!",
        "C": "Needs improvement.",
        "D": "Work harder!",
        "F": "Failed. Please focus on weak areas."
    }
    return comments.get(grade, "No comment")

# Process grading for students
def process_grading_grading(df):
    try:
        # Ensure required columns are present
        required_columns = ["Name", "Subject", "Score", "Maximum Marks"]
        for col in required_columns:
            if col not in df.columns:
                raise ValueError(f"Missing required column: {col}")

        # Group data by student and subject
        results = []
        for name, group in df.groupby("Name"):
            for _, row in group.iterrows():
                percentage = (row["Score"] / row["Maximum Marks"]) * 100
                grade = calculate_grade_grading(percentage)
                comment = get_comment_grading(grade)

                results.append({
                    "Name": name,
                    "Subject": row["Subject"],
                    "Score": row["Score"],
                    "Maximum Marks": row["Maximum Marks"],
                    "Percentage": round(percentage, 2),
                    "Grade": grade,
                    "Comment": comment
                })

        return pd.DataFrame(results)

    except Exception as e:
        st.error(f"Error processing grading: {e}")
        return None

# Streamlit app
def grading():
    st.title("Automated Grading System with Subject Details")
    st.write("Upload an Excel or CSV file containing student data with maximum marks and subjects to calculate grades.")

    # File uploader
    uploaded_file = st.file_uploader("Upload your Excel or CSV file", type=["csv", "xlsx"])

    if uploaded_file:
        try:
            # Read the uploaded file
            if uploaded_file.name.endswith(".xlsx"):
                df = pd.read_excel(uploaded_file)
            else:
                df = pd.read_csv(uploaded_file)

            # Display the uploaded data
            st.subheader("Uploaded Data")
            st.write(df)

            # Process grading
            st.subheader("Grading Results")
            results = process_grading_grading(df)
            if results is not None:
                st.write(results)

                # Option to download the results
                st.subheader("Download Results")
                output_file = "grading_results_with_subjects.xlsx"
                results.to_excel(output_file, index=False)

                with open(output_file, "rb") as file:
                    st.download_button(
                        label="Download Excel",
                        data=file,
                        file_name=output_file,
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )

        except Exception as e:
            st.error(f"Error processing the file: {e}")




import pandas as pd
import streamlit as st
import matplotlib.pyplot as plt
from io import BytesIO

# Plot a graph for subject vs. percentage for a chosen student
def plot_subject_vs_percentage(subject_data, student_name):
    student_data = subject_data[subject_data["Name"] == student_name]
    if student_data.empty:
        st.warning(f"No data available for {student_name}.")
        return None

    plt.figure(figsize=(8, 5))
    plt.bar(student_data["Subject"], student_data["Percentage"], color="skyblue", edgecolor="black")
    plt.xlabel("Subject")
    plt.ylabel("Percentage")
    plt.title(f"Performance of {student_name}")
    plt.ylim(0, 100)
    plt.xticks(rotation=45)
    buf = BytesIO()
    plt.savefig(buf, format="png")
    buf.seek(0)
    plt.clf()
    return buf

# Plot a comparative graph for subject vs. average percentage across all students
def plot_subject_vs_avg_percentage(subject_data):
    avg_percentage = subject_data.groupby("Subject")["Percentage"].mean()
    if avg_percentage.empty:
        st.warning("No data available for subjects.")
        return None

    plt.figure(figsize=(8, 5))
    plt.bar(avg_percentage.index, avg_percentage.values, color="lightgreen", edgecolor="black")
    plt.xlabel("Subject")
    plt.ylabel("Average Percentage")
    plt.title("Average Performance per Subject")
    plt.ylim(0, 100)
    plt.xticks(rotation=45)
    buf = BytesIO()
    plt.savefig(buf, format="png")
    buf.seek(0)
    plt.clf()
    return buf

# Streamlit app for performance graph
def performance_graph():
    st.title("Performance Graph Module")
    st.write("Upload the Excel or CSV file generated by the Grading module to visualize performance and analyze data.")

    # File uploader
    uploaded_file = st.file_uploader("Upload the file generated by the Grading module", type=["csv", "xlsx"])

    if uploaded_file:
        try:
            # Read the uploaded file
            if uploaded_file.name.endswith(".xlsx"):
                df = pd.read_excel(uploaded_file)
            else:
                df = pd.read_csv(uploaded_file)

            # Display the uploaded data
            st.subheader("Uploaded Data")
            st.write(df)

            # Check required columns
            required_columns = ["Name", "Subject", "Score", "Maximum Marks", "Percentage", "Grade"]
            for col in required_columns:
                if col not in df.columns:
                    st.error(f"Missing required column: {col}. Ensure the file is generated by the Grading module.")
                    return

            # Graph dropdown selection
            st.subheader("Visualizations")
            graph_option = st.selectbox("Select a Graph Type", ["Individual Student Performance", "Average Performance per Subject"])

            if graph_option == "Individual Student Performance":
                selected_student = st.selectbox("Select a student to view their performance:", df["Name"].unique())
                if selected_student:
                    buf = plot_subject_vs_percentage(df, selected_student)
                    if buf:
                        st.image(buf, caption=f"Performance of {selected_student}", use_column_width=True)
                        st.download_button(
                            label="Download Graph (Individual Student)",
                            data=buf,
                            file_name=f"{selected_student}_performance.png",
                            mime="image/png"
                        )

            elif graph_option == "Average Performance per Subject":
                buf_avg = plot_subject_vs_avg_percentage(df)
                if buf_avg:
                    st.image(buf_avg, caption="Average Performance per Subject", use_column_width=True)
                    st.download_button(
                        label="Download Graph (All Students)",
                        data=buf_avg,
                        file_name="average_subject_performance.png",
                        mime="image/png"
                    )

        except Exception as e:
            st.error(f"Error processing the file: {e}")


import streamlit as st
from docx import Document
from io import BytesIO
import openai

# Function to generate a paragraph using OpenAI
def generate_paragraph(prompt):
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are an educational content creator."},
                {"role": "user", "content": prompt},
            ],
        )
        return response["choices"][0]["message"]["content"]
    except Exception as e:
        st.error(f"Error generating paragraph: {e}")
        return None

# Function to generate questions using OpenAI
def paragraph_questions(paragraph, question_prompt, medium):
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": f"You are an educational content creator creating questions in {medium}."},
                {"role": "user", "content": f"Here is the paragraph: {paragraph}"},
                {"role": "user", "content": question_prompt},
            ],
        )
        return response["choices"][0]["message"]["content"]
    except Exception as e:
        st.error(f"Error generating questions: {e}")
        return None

# Function to create a DOCX file
def create_docx(content, filename="questions.docx"):
    doc = Document()
    for line in content.split("\n"):
        doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit App
def Paragraph_Based_Questions():
    st.title("Paragraph and Question Generator")
    st.write("Generate a paragraph and corresponding questions based on input parameters.")

    # Inputs for question generation
    boards = ["CBSE", "ICSE", "State Board", "IB", "Others"]
    classes = [f"Class {i}" for i in range(1, 13)]
    subjects = ["Mathematics", "Science", "English", "History", "Geography", "Others"]
    question_types = ["True/False", "Yes/No", "Short Answers", "Very Short Answers", "Mixed"]
    mediums = ["English", "Hindi", "Others"]

    board = st.selectbox("Select Board", boards)
    if board == "Others":
        board = st.text_input("Enter Board Name:")

    grade = st.selectbox("Select Class", classes)
    subject = st.selectbox("Select Subject", subjects)
    if subject == "Others":
        subject = st.text_input("Enter Subject Name:")

    topic = st.text_input("Enter Topic", placeholder="E.g., Photosynthesis, Algebra")
    question_type = st.selectbox("Select Question Type", question_types)
    medium = st.selectbox("Select Medium of Instruction", mediums)
    if medium == "Others":
        medium = st.text_input("Enter Medium:")

    num_questions = st.number_input("Enter Number of Questions", min_value=1, max_value=50, step=1)
    max_marks = st.number_input("Enter Maximum Marks (Optional)", min_value=1, step=1, value=1)

    include_answers = st.radio("Include Answers in the Questions?", ["Yes", "No"]) == "Yes"

    # Generate paragraph and questions button
    if st.button("Generate Paragraph and Questions"):
        if not all([board, grade, subject, topic, question_type, medium]):
            st.warning("Please fill in all the required fields.")
        else:
            # Prepare the paragraph prompt
            paragraph_prompt = f"""
            Generate an educational paragraph for {subject} for {grade} under the {board} board.
            The topic is "{topic}" and the medium of instruction is {medium}.
            """

            with st.spinner("Generating paragraph..."):
                paragraph = generate_paragraph(paragraph_prompt)

            if paragraph:
                st.subheader("Generated Paragraph")
                st.write(paragraph)

                # Prepare the question prompt
                question_prompt = f"""
                Based on the provided paragraph, generate {num_questions} {question_type} questions in {medium}.
                {"Include answers for each question." if include_answers else "Do not include answers."}
                {"The total marks should not exceed " + str(max_marks) + "." if max_marks > 0 else ""}
                """

                with st.spinner("Generating questions..."):
                    questions = paragraph_questions(paragraph, question_prompt, medium)

                if questions:
                    st.subheader("Generated Questions")
                    st.write(questions)

                    # Allow download as DOCX
                    combined_content = f"Paragraph:\n\n{paragraph}\n\nQuestions:\n\n{questions}"
                    docx_file = create_docx(combined_content)
                    st.download_button(
                        label="Download Paragraph and Questions as DOCX",
                        data=docx_file,
                        file_name="paragraph_and_questions.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )







# Main function to handle page navigation
def main():
    if st.session_state['page'] == 'home':
        landing_page()
    elif st.session_state['page'] == 'login':
        login_page()
    elif st.session_state['page'] == 'main' and st.session_state['logged_in']:
        main_app()
    else:
        st.error("An error occurred. Please refresh the page.")

def call_openai_api(prompt, max_tokens=1500):
    """
    Calls OpenAI API with the given prompt and returns the generated content.
    """
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=max_tokens
    )
    return response["choices"][0]["message"]["content"]


if __name__ == "__main__":
    main()


